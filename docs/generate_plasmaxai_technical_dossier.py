from __future__ import annotations

import json
from pathlib import Path

import joblib
import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
RESEARCH_DIR = ROOT / "research"
DOCS_DIR = ROOT / "docs"
OUT_DOC = DOCS_DIR / "PlasmaXAI_Technical_Research_Architecture_Dossier.docx"

NOVEL_DIR = RESEARCH_DIR / "outputs" / "novel"
FIG_DIR = NOVEL_DIR / "figures"
PACKAGE_FIG_DIR = RESEARCH_DIR / "package" / "diagrams_and_images" / "paper_figures"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    if "EquationBlock" not in doc.styles:
        style = doc.styles.add_style("EquationBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)

    if "CaptionSmall" not in doc.styles:
        style = doc.styles.add_style("CaptionSmall", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(9.5)
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(8)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraphs(doc: Document, text: str) -> None:
    for block in [chunk.strip() for chunk in text.strip().split("\n\n") if chunk.strip()]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.add_run(block)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_equation(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="EquationBlock")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        set_cell_shading(hdr_cells[idx], "1F2937")
        for paragraph in hdr_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def add_figure(doc: Document, image_path: Path, caption: str, technical_note: str, width: float = 6.2) -> None:
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(style="CaptionSmall")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    r.bold = True
    add_paragraphs(doc, technical_note)


def load_context() -> dict[str, object]:
    summary = json.loads((NOVEL_DIR / "novel_summary.json").read_text(encoding="utf-8"))
    bootstrap = json.loads((NOVEL_DIR / "bootstrap_summary.json").read_text(encoding="utf-8"))
    comparison_df = pd.read_csv(NOVEL_DIR / "plasmaxai_extended_model_comparison.csv")
    ablation_df = pd.read_csv(NOVEL_DIR / "plasmaxai_ablation_results.csv")
    patient_df = pd.read_csv(NOVEL_DIR / "patient_counterfactual_summary.csv")
    feature_bundle = joblib.load(NOVEL_DIR / "fusion_scalers.joblib")
    feature_blocks = feature_bundle["feature_blocks"]

    return {
        "summary": summary,
        "bootstrap": bootstrap,
        "comparison_df": comparison_df,
        "ablation_df": ablation_df,
        "patient_df": patient_df,
        "feature_dims": {
            "resnet": len(feature_blocks[0]),
            "densenet": len(feature_blocks[1]),
            "morphology": len(feature_blocks[2]),
            "counterfactual": len(feature_blocks[3]),
            "score": len(feature_blocks[4]),
        },
    }


def build_document() -> Document:
    ctx = load_context()
    summary = ctx["summary"]
    bootstrap = ctx["bootstrap"]
    comparison_df: pd.DataFrame = ctx["comparison_df"]
    ablation_df: pd.DataFrame = ctx["ablation_df"]
    patient_df: pd.DataFrame = ctx["patient_df"]
    dims = ctx["feature_dims"]

    plasmaxai_row = comparison_df.loc[comparison_df["model"] == "PlasmaXAI"].iloc[0]
    best_baseline = comparison_df.loc[comparison_df["model"] != "PlasmaXAI"].sort_values("accuracy", ascending=False).iloc[0]
    patient_summary = summary["patient_summary"]
    diseased_count = int((patient_df["diagnosis"] == "diseased").sum())
    normal_count = int((patient_df["diagnosis"] == "normal").sum())

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "PlasmaXAI Technical Research and Architecture Dossier"
    doc.core_properties.author = "Arnob Aich Anurag / OpenAI Codex"

    add_title(
        doc,
        "PlasmaXAI Technical Research and Architecture Dossier",
        "Equation-driven, layer-resolved, end-to-end technical description of the PlasmaXAI research framework and deployment system",
    )

    add_heading(doc, "1. Research Scope and Technical Intent", 1)
    add_paragraphs(
        doc,
        f"""PlasmaXAI was developed to solve a specific computational pathology problem: reliable recognition of malignant plasma cells from microscopy imagery while preserving enough interpretability for doctor-facing review. The system therefore had to satisfy two constraints simultaneously. First, it had to outperform simpler baselines on cell-level classification. Second, its internal evidence path had to be decomposable into clinically readable cues rather than ending at a single opaque probability.

The final research framework is not a single convolutional network trained in isolation. It is a multi-branch system that fuses two deep image backbones, a handcrafted morphology vector, a counterfactual boundary vector, and a compact score branch. The core novelty is that counterfactual information is not applied after prediction as a post-hoc add-on. Instead, the counterfactual branch directly modulates the latent image branches and influences the learned modality weights used during fusion.

At the research level, the framework achieved {summary['novel_fusion_test_metrics']['accuracy']:.2f}% accuracy, {summary['novel_fusion_test_metrics']['weighted_f1']:.2f}% weighted F1, {summary['novel_fusion_test_metrics']['plasma_precision']:.2f}% plasma precision, {summary['novel_fusion_test_metrics']['plasma_recall']:.2f}% plasma recall, and {summary['novel_fusion_test_metrics']['auc']:.2f}% AUC at the main novel-fusion operating point. At the deployment-oriented comparison operating point stored in the current benchmark file, PlasmaXAI reaches {plasmaxai_row['accuracy']:.2f}% accuracy, {plasmaxai_row['plasma_precision']:.2f}% plasma precision, {plasmaxai_row['plasma_recall']:.2f}% plasma recall, and {plasmaxai_row['auc']:.2f}% AUC."""
    )

    add_heading(doc, "2. Data Assets, Problem Framing, and Split Protocol", 1)
    add_paragraphs(
        doc,
        f"""The model-development split is based on PCMMD Set 1, which is used for train/validation/test development. A second patient-organized subset (PCMMD Set 2) is used for exploratory patient-level signature analysis. The research problem is binary classification at the cropped-cell level: class 1 corresponds to plasma cells and class 0 corresponds to non-plasma cells. The pipeline then extends this cell-level view into patient-level aggregation.

The exact main split is stratified 70/15/15. In implementation terms, the full dataframe is first divided into 70% train and 30% temporary data, then the temporary portion is split again into 15% validation and 15% test. This preserves the class ratio in every split and avoids threshold selection on the held-out test set.

The observed split counts in the validated run are 2482 training images, 532 validation images, and 532 test images. For patient-level analysis, {diseased_count + normal_count} patient entities were evaluated from the detection/patients hierarchy, with {diseased_count} diseased and {normal_count} normal patients in the saved patient summary artifact."""
    )
    add_table(
        doc,
        ["Data subset", "Purpose", "Observed size", "Notes"],
        [
            ["PCMMD Set 1 - train", "Parameter learning", "2482", "Stratified cell-level split"],
            ["PCMMD Set 1 - validation", "Threshold/model selection", "532", "Used for operating-point search"],
            ["PCMMD Set 1 - test", "Final hold-out evaluation", "532", "Never used for threshold search"],
            ["PCMMD Set 2 - patients", "Patient-level signature study", str(diseased_count + normal_count), "YOLO-labeled patient folders"],
        ],
    )

    add_heading(doc, "3. Hardware and Software Runtime", 1)
    add_paragraphs(
        doc,
        """The validated research run was executed with CUDA enabled. The verified GPU in the project record is an NVIDIA GeForce RTX 3060 with 12288 MiB of VRAM, and the CUDA runtime reported by the local environment was CUDA 13.1. The Python version verified in the same environment was Python 3.13.2.

The technical stack spans both research and deployment layers. The model training and inference code rely on PyTorch, torchvision, timm, OpenCV, scikit-learn, pandas, and NumPy. The doctor-facing website is implemented in Next.js 16.2.2 with React 19.2.4, while the remote inference service is implemented with FastAPI. Artifact generation for the current request uses python-docx."""
    )
    add_table(
        doc,
        ["Layer", "Main tools", "Role in the project"],
        [
            ["Deep learning", "PyTorch, torchvision, timm", "Backbone loading, fusion model, CUDA execution"],
            ["Classical ML", "scikit-learn", "Logistic counterfactual model, scaling, metrics"],
            ["Image processing", "OpenCV, PIL", "Morphology extraction, byte decoding, rendering"],
            ["Research output", "pandas, NumPy, matplotlib, seaborn", "Metrics, CSVs, and diagrams"],
            ["Web platform", "Next.js, React, Bun", "Doctor workspace and reporting UI"],
            ["Inference API", "FastAPI", "Remote analysis endpoint for case review"],
        ],
    )

    add_heading(doc, "4. End-to-End Research Pipeline", 1)
    add_paragraphs(
        doc,
        """The full research pipeline can be decomposed into six stages. Stage 1 loads microscopy crops and labels from the cell-level dataset. Stage 2 derives two independent deep image embeddings using ResNet50 and DenseNet121 checkpoints that had already been optimized in the earlier hybrid stage. Stage 3 extracts a handcrafted morphology vector from every image. Stage 4 fits a logistic counterfactual model on morphology and transforms every sample into a counterfactual descriptor relative to the benign boundary. Stage 5 feeds all five blocks into the learned fusion network. Stage 6 aggregates cell-level outputs into patient-level summary statistics and generates quantitative figures.

This ordering matters. The framework does not ask the fusion classifier to infer everything from RGB alone. Instead, it presents the final classifier with complementary signals: texture- and structure-sensitive deep image features, explicit cell morphology, and a mathematically constructed estimate of how far each cell would have to move in feature space to become benign-leaning."""
    )

    add_heading(doc, "5. Formal Input Definition", 1)
    add_paragraphs(
        doc,
        f"""For each microscopy crop I, the system constructs five input blocks:

1. ResNet embedding e_r in R^{dims['resnet']}
2. DenseNet embedding e_d in R^{dims['densenet']}
3. Morphology vector x_m in R^{dims['morphology']}
4. Counterfactual descriptor x_cf in R^{dims['counterfactual']}
5. Compact score vector x_s in R^{dims['score']}

The score vector is not arbitrary. It contains the ResNet plasma probability, the DenseNet plasma probability, and the morphology-logistic counterfactual plasma probability. Therefore, the fusion model sees both latent features and calibrated branch-level scalar evidence."""
    )
    add_equation(doc, "(1) x_s = [ p_resnet , p_densenet , p_cf ]")

    add_heading(doc, "6. Morphology Feature Engineering", 1)
    add_paragraphs(
        doc,
        """The morphology branch extracts low-dimensional descriptors intended to preserve clinically meaningful cell geometry and stain behavior. These features are computed directly from the microscopy crop after image decoding and basic mask construction. The vector has 10 dimensions and is deliberately compact so that every coordinate can be interpreted by a reviewer or traced back to a measurable image quantity.

The implemented morphology feature names are mean_r, mean_g, mean_b, stain_intensity, nucleus_area, cell_area, cytoplasm_area, nc_ratio, granularity, and roundness. The exact code estimates these terms from grayscale intensity, thresholded regions, and contour geometry. Conceptually, the branch is encoding color concentration, nucleus-to-cytoplasm balance, texture heterogeneity, and gross cell shape."""
    )
    add_equation(doc, "(2) mu_R = (1 / |Omega|) * sum_{(u,v) in Omega} I_R(u,v)")
    add_equation(doc, "(3) mu_G = (1 / |Omega|) * sum_{(u,v) in Omega} I_G(u,v)")
    add_equation(doc, "(4) mu_B = (1 / |Omega|) * sum_{(u,v) in Omega} I_B(u,v)")
    add_equation(doc, "(5) I_stain = 255 - (1 / |Omega|) * sum_{(u,v) in Omega} Gray(u,v)")
    add_equation(doc, "(6) A_nuc = |Omega_nucleus|,   A_cell = |Omega_cell|,   A_cyto = max(A_cell - A_nuc, 0)")
    add_equation(doc, "(7) r_NC = A_nuc / (A_cyto + eps)")
    add_equation(doc, "(8) G_tex = std(Gray(Omega_cell))")
    add_equation(doc, "(9) R_shape = 4 * pi * A_cell / (P_cell^2 + eps)")
    add_bullets(
        doc,
        [
            "mean_r, mean_g, and mean_b quantify channel-wise staining behavior and later reappear as the dominant patient-level counterfactual drivers in the saved patient summary.",
            "stain_intensity captures whether the cell occupies a darker and more saturated field than surrounding benign-like crops.",
            "nucleus_area, cytoplasm_area, and nc_ratio encode the nuclear burden of the crop, which is central to plasma-cell morphology review.",
            "granularity acts as a texture dispersion proxy and lets the model distinguish smooth benign-looking regions from more heterogeneous cell interiors.",
            "roundness is a coarse structural stabilizer so the tabular stream does not rely solely on stain statistics.",
        ],
    )
    add_table(
        doc,
        ["Morphology feature", "Mathematical source", "Clinical interpretation"],
        [
            ["mean_r / mean_g / mean_b", "Channel means over cell support", "Relative stain and color concentration"],
            ["stain_intensity", "Inverted grayscale mean", "Overall darkness / stain load"],
            ["nucleus_area", "Pixel count in nucleus mask", "Estimated nuclear burden"],
            ["cell_area", "Pixel count in full cell mask", "Global cell size"],
            ["cytoplasm_area", "cell_area - nucleus_area", "Available cytoplasmic volume"],
            ["nc_ratio", "nucleus_area / cytoplasm_area", "Nucleus-to-cytoplasm emphasis"],
            ["granularity", "Std. dev. over grayscale cell region", "Internal texture heterogeneity"],
            ["roundness", "4*pi*A / P^2", "Contour compactness"],
        ],
    )

    add_heading(doc, "7. Deep Backbone Feature Extraction", 1)
    add_paragraphs(
        doc,
        f"""The image pathway is intentionally redundant. PlasmaXAI uses two separately trained backbone families because they induce different feature geometries. The ResNet50 branch outputs an embedding e_r in R^{dims['resnet']}, while the DenseNet121 branch outputs e_d in R^{dims['densenet']}. These embeddings were inherited from the earlier optimized search stage and then reused inside the novel fusion framework.

ResNet50 contributes stronger stage-wise residual semantics and typically captures coarse structural organization. DenseNet121 preserves dense feature reuse and often responds more strongly to local texture continuity and stain transitions. The reason for keeping both is that plasma-cell microscopy contains both macro-structure and subtle local morphology. A single backbone can miss one of those regimes."""
    )
    add_equation(doc, "(10) e_r = f_resnet(I),   e_d = f_densenet(I)")
    add_equation(doc, "(11) p_resnet = sigma(w_r^T e_r + b_r),   p_densenet = sigma(w_d^T e_d + b_d)")
    add_paragraphs(
        doc,
        """The scalar branch probabilities in Equation (11) are not the final decision. They are retained because they serve as compact evidence values in the score block. This gives the later fusion layer access to both full latent embeddings and already-compressed probabilistic opinions from the image branches."""
    )
    add_table(
        doc,
        ["Image branch", "Embedding dimension", "Functional role inside PlasmaXAI"],
        [
            ["ResNet50", str(dims["resnet"]), "Residual feature hierarchy; coarse morphology and structural contrast"],
            ["DenseNet121", str(dims["densenet"]), "Dense feature reuse; fine texture and stain-transition behavior"],
            ["Score block", str(dims["score"]), "Compact branch-level scalar evidence for later fusion"],
        ],
    )

    add_heading(doc, "8. Counterfactual Boundary Model", 1)
    add_paragraphs(
        doc,
        f"""The counterfactual branch is what differentiates PlasmaXAI from a standard late-fusion ensemble. A logistic model is fit on the standardized morphology features. Let x_m be the raw morphology vector and z its standardized version. The logistic decision margin m = w^T z + b becomes the starting point for a geometry-aware estimate of how the sample would need to move in morphology space to become benign-leaning.

The current saved counterfactual descriptor has dimension {dims['counterfactual']}. It contains both absolute position signals and directional shift signals, including the signed standardized features, the minimal boundary correction, the prototype shift toward the benign centroid, the scalar logistic probability, and the margin distance. This block is not just an explanation summary. It is one of the actual trainable inputs to the fusion classifier."""
    )
    add_equation(doc, "(12) z = StandardScaler(x_m)")
    add_equation(doc, "(13) m = w^T z + b,   p_cf = sigma(m)")
    add_equation(doc, "(14) delta_boundary = - ( max(m, 0) / ( ||w||_2^2 + eps ) ) * w")
    add_equation(doc, "(15) delta_proto = mu_benign - z")
    add_equation(doc, "(16) d_cf = max(m, 0) / ( ||w||_2 + eps )")
    add_paragraphs(
        doc,
        """Equation (14) is the shortest-step correction that projects a plasma-leaning standardized morphology vector back toward the logistic boundary. Equation (15) complements that by measuring the shift from the current sample toward the empirical benign prototype. Together, the descriptor captures both local boundary geometry and cohort-level benign direction.

This is the key conceptual move of PlasmaXAI: the counterfactual branch is model-native. It is computed before fusion, scaled, embedded, and allowed to influence the latent image branches themselves through learned gates."""
    )

    add_heading(doc, "9. Counterfactual-Guided Fusion Network", 1)
    add_paragraphs(
        doc,
        f"""The final learned model is CounterfactualGuidedFusionNet. Its saved best configuration is hidden_dim = {summary['best_fusion_config']['hidden_dim']}, tabular_dim = {summary['best_fusion_config']['tabular_dim']}, dropout = {summary['best_fusion_config']['dropout']}, learning rate = {summary['best_fusion_config']['lr']}, weight decay = {summary['best_fusion_config']['weight_decay']}, batch size = {summary['best_fusion_config']['batch_size']}, patience = {summary['best_fusion_config']['patience']}, and maximum epochs = {summary['best_fusion_config']['epochs']}. The model does not concatenate raw blocks directly. Instead, each block is first normalized and projected into a controlled latent dimension.

Let h = 256 and t = 128 for the best saved configuration. The image branches are projected into h-dimensional latent spaces, while the morphology, counterfactual, and score branches are projected into t-dimensional latent spaces. The counterfactual embedding then emits two multiplicative gates, one for the ResNet latent and one for the DenseNet latent. This means that counterfactual evidence changes how much of each latent image coordinate survives into the fusion stage."""
    )
    add_equation(doc, "(17) r_hat = Dropout( GELU( W_r * LN(e_r) + b_r ) )")
    add_equation(doc, "(18) d_hat = Dropout( GELU( W_d * LN(e_d) + b_d ) )")
    add_equation(doc, "(19) m_hat = Dropout( GELU( W_m * LN(x_m) + b_m ) )")
    add_equation(doc, "(20) c_hat = Dropout( GELU( W_c * LN(x_cf) + b_c ) )")
    add_equation(doc, "(21) s_hat = Dropout( GELU( W_s * LN(x_s) + b_s ) )")
    add_equation(doc, "(22) g_r = sigmoid( W_cr * c_hat + b_cr ),   g_d = sigmoid( W_cd * c_hat + b_cd )")
    add_equation(doc, "(23) r_tilde = r_hat odot g_r,   d_tilde = d_hat odot g_d")
    add_equation(doc, "(24) s_tilde = Dropout( GELU( W_ms [m_hat ; c_hat] + b_ms ) ) + s_hat")
    add_equation(doc, "(25) alpha = softmax( W_gate * LN([r_tilde ; d_tilde ; m_hat ; c_hat ; s_tilde]) + b_gate )")
    add_equation(doc, "(26) u = [ alpha_1 U_r(r_tilde) + alpha_2 U_d(d_tilde) + alpha_3 U_m(m_hat) + alpha_4 U_c(c_hat) + U_s(s_tilde) ; alpha_1 U_r(r_tilde) + alpha_2 U_d(d_tilde) ]")
    add_equation(doc, "(27) y = W_o * Dropout( GELU( W_h * LN(u) + b_h ) ) + b_o")
    add_equation(doc, "(28) p = softmax(y)")
    add_table(
        doc,
        ["Module", "Exact implementation", "Purpose"],
        [
            ["Res encoder", "LayerNorm(2048) -> Linear(2048,256) -> GELU -> Dropout", "Project ResNet latent into shared hidden scale"],
            ["Dense encoder", "LayerNorm(1024) -> Linear(1024,256) -> GELU -> Dropout", "Project DenseNet latent into shared hidden scale"],
            ["Morph encoder", "LayerNorm(10) -> Linear(10,128) -> GELU -> Dropout", "Compress interpretable morphology features"],
            ["CF encoder", f"LayerNorm({dims['counterfactual']}) -> Linear({dims['counterfactual']},128) -> GELU -> Dropout", "Encode counterfactual geometry"],
            ["Score encoder", "LayerNorm(3) -> Linear(3,128) -> GELU -> Dropout", "Encode compact branch probabilities"],
            ["CF-to-image gates", "Linear(128,256) + Sigmoid twice", "Modulate each image branch with counterfactual evidence"],
            ["Morph-to-score refiner", "Linear(256,128) -> GELU -> Dropout", "Inject morphology+CF context into score latent"],
            ["Modality gate", "LayerNorm(896) -> Linear(896,4) -> softmax", "Learn relative importance of res/dense/morph/cf streams"],
            ["Classifier", "LayerNorm(512) -> Linear(512,256) -> GELU -> Dropout -> Linear(256,2)", "Produce final class logits"],
        ],
    )
    add_paragraphs(
        doc,
        """A critical technical point is that the score stream is not gated away by the four-way modality softmax. Instead, the score latent is always carried into the final concatenation after unification, while the four-way gate controls the relative contributions of the two image latents, morphology latent, and counterfactual latent. This makes the architecture both stable and expressive: branch probabilities are always available, but the richer latent streams can be emphasized or suppressed on a per-sample basis."""
    )

    add_heading(doc, "10. Training Objective, Optimization, and Calibration", 1)
    add_paragraphs(
        doc,
        """The fusion network is trained as a weighted binary classifier on the development split. Because the malignant/plasma decision is the clinically sensitive class, the project tracks plasma recall explicitly during model search and threshold selection. The training loss is a class-weighted cross-entropy over the final logits y, while the optimization loop separately monitors validation recall and validation AUC to choose acceptable operating points.

The final model-selection logic in the saved research artifacts uses a validation threshold search rather than a fixed 0.5 cutoff. This is an important part of the performance story: the raw network learns the decision surface, but the chosen operating point determines the precision-recall tradeoff seen during deployment benchmarking."""
    )
    add_equation(doc, "(29) L_CE = - sum_{i=1}^{N} sum_{k in {0,1}} w_k * 1[y_i = k] * log p_{ik}")
    add_equation(doc, "(30) theta* = argmax_theta  M_val(theta)")
    add_equation(doc, "(31) M_val(theta) = lambda_1 * Recall_val(theta) + lambda_2 * AUC_val(theta) + lambda_3 * Acc_val(theta)")
    add_paragraphs(
        doc,
        f"""For the main novel-fusion research result, the saved threshold is {summary['novel_threshold']:.2f}. For the doctor-facing benchmark file used later during deployment-oriented comparison, the selected threshold is {plasmaxai_row['threshold']:.2f}. This distinction is why multiple PlasmaXAI performance summaries appear in the project history. They correspond to different operating points on the same learned system rather than different architectures."""
    )
    add_table(
        doc,
        ["Hyperparameter", "Saved best value", "Technical effect"],
        [
            ["hidden_dim", str(summary["best_fusion_config"]["hidden_dim"]), "Latent width for image branches and classifier bottleneck"],
            ["tabular_dim", str(summary["best_fusion_config"]["tabular_dim"]), "Latent width for morphology/counterfactual/score branches"],
            ["dropout", str(summary["best_fusion_config"]["dropout"]), "Regularization across all encoders and classifier layers"],
            ["learning_rate", str(summary["best_fusion_config"]["lr"]), "Adam step size"],
            ["weight_decay", str(summary["best_fusion_config"]["weight_decay"]), "L2-style regularization"],
            ["batch_size", str(summary["best_fusion_config"]["batch_size"]), "GPU batch throughput"],
            ["patience", str(summary["best_fusion_config"]["patience"]), "Early stopping tolerance"],
            ["epochs", str(summary["best_fusion_config"]["epochs"]), "Maximum allowed training horizon"],
        ],
    )

    add_heading(doc, "11. Ablation, Bootstrap, and Patient Aggregation", 1)
    add_paragraphs(
        doc,
        """Three validation layers were used to test whether the framework's gains were real rather than accidental. First, an ablation study removed either morphology or the counterfactual path. Second, bootstrap resampling estimated confidence intervals for the headline metrics. Third, an exploratory patient-level aggregation study tested whether the cell-level outputs formed coherent patient signatures."""
    )
    add_table(
        doc,
        ["Ablation setting", "Accuracy", "Weighted F1", "Plasma precision", "Plasma recall", "AUC"],
        [
            [
                row["label"],
                f"{row['test_accuracy']:.2f}",
                f"{row['test_weighted_f1']:.2f}",
                f"{row['test_plasma_precision']:.2f}",
                f"{row['test_plasma_recall']:.2f}",
                f"{row['test_auc']:.2f}",
            ]
            for _, row in ablation_df.iterrows()
        ],
    )
    add_paragraphs(
        doc,
        f"""The strongest ablation signal is the drop caused by removing the counterfactual path. In the saved ablation CSV, plasma recall falls to {ablation_df.loc[ablation_df['label'] == 'Without Counterfactual Path', 'test_plasma_recall'].iloc[0]:.2f}% even though AUC remains high. This is exactly what one would expect if the counterfactual path is helping the model commit to plasma-positive decisions in hard boundary cases rather than simply improving separability in a generic sense.

Bootstrap validation gives additional stability evidence. The novel framework's bootstrap mean accuracy is {bootstrap['novel']['accuracy']['mean']:.2f}% with a 95% interval of [{bootstrap['novel']['accuracy']['low95']:.2f}, {bootstrap['novel']['accuracy']['high95']:.2f}] and bootstrap mean AUC is {bootstrap['novel']['auc']['mean']:.2f}% with a 95% interval of [{bootstrap['novel']['auc']['low95']:.2f}, {bootstrap['novel']['auc']['high95']:.2f}]."""
    )
    add_equation(doc, "(32) PatientMeanProb_j = (1 / n_j) * sum_{i=1}^{n_j} p_i")
    add_equation(doc, "(33) PatientMeanCF_j = (1 / n_j^+) * sum_{i : plasma_i = 1} d_cf,i")
    add_equation(doc, "(34) Consistency_j = mean( cosine(Delta_i, Delta_k) ) for i != k")
    add_equation(doc, "(35) Score_patient = 0.6 * PatientMeanProb + 0.4 * PatientMeanCF")
    add_paragraphs(
        doc,
        f"""At patient level, the saved summary reports patient-level AUC = {patient_summary['patient_auc_mean_novel_prob']:.3f} when using mean novel probability and AUC = {patient_summary['patient_auc_combined_signature']:.3f} when using the combined score in Equation (35). The diseased-vs-normal group means also indicate that diseased patients had higher mean novel probability ({patient_summary['group_means']['mean_novel_prob']['diseased']:.4f} vs {patient_summary['group_means']['mean_novel_prob']['normal']:.4f}) and lower mean counterfactual distance to the benign boundary ({patient_summary['group_means']['mean_cf_distance_plasma']['diseased']:.4f} vs {patient_summary['group_means']['mean_cf_distance_plasma']['normal']:.4f})."""
    )

    add_heading(doc, "12. Quantitative Result Summary", 1)
    add_paragraphs(
        doc,
        f"""Relative to the best non-PlasmaXAI baseline in the current comparison file ({best_baseline['model']}), PlasmaXAI is competitive or better across the clinically important metrics while preserving interpretability. The benchmark table below is taken from the currently saved comparison artifact used by the project website and deployment logic."""
    )
    top_rows = comparison_df[["model", "accuracy", "weighted_f1", "plasma_precision", "plasma_recall", "auc"]].copy()
    add_table(
        doc,
        ["Model", "Accuracy", "Weighted F1", "Plasma precision", "Plasma recall", "AUC"],
        [
            [
                row["model"],
                f"{row['accuracy']:.2f}",
                f"{row['weighted_f1']:.2f}",
                f"{row['plasma_precision']:.2f}",
                f"{row['plasma_recall']:.2f}",
                f"{row['auc']:.2f}",
            ]
            for _, row in top_rows.iterrows()
        ],
    )
    add_paragraphs(
        doc,
        """The two performance snapshots used in the project should be interpreted carefully. The novel_summary.json values describe the main novel-fusion research run after its own validation threshold search. The plasmaxai_extended_model_comparison.csv values describe the later deployment-oriented operating point. Both are valid, but they answer slightly different operational questions."""
    )

    add_heading(doc, "13. Research-to-Product Integration", 1)
    add_paragraphs(
        doc,
        """The deployed system mirrors the research decomposition. The doctor-facing interface is implemented in Next.js and presents cases, images, focus maps, explainability diagrams, editable checklists, and generated reports. The inference server is implemented in FastAPI and exposes a POST /cases endpoint. At inference time, the server loads the saved operating point, scalers, counterfactual bundle, backbone checkpoints, and final fusion weights.

The data path is therefore: doctor upload -> storage reference -> inference request -> image decoding -> morphology extraction -> counterfactual descriptor construction -> backbone feature extraction -> learned fusion -> clinical interpretation and report generation -> database persistence. In the Supabase schema, this path is reflected by the tables organizations, profiles, patients, cases, case_images, predictions, explanations, reports, review_requests, and audit_logs. The website also contains a local fallback persistence layer for offline/demo execution, but the production architecture is explicitly organized for web + inference + database separation."""
    )
    add_table(
        doc,
        ["Deployment layer", "Concrete implementation", "Technical responsibility"],
        [
            ["Frontend", "apps/web (Next.js, React, Bun)", "Doctor dashboard, case review, patient CRUD, report editing"],
            ["Inference API", "apps/inference (FastAPI)", "Load model assets and return structured predictions"],
            ["Model assets", "artifacts/models + research/outputs/novel", "Checkpoints, scalers, counterfactual bundles"],
            ["Database", "Supabase Postgres schema", "Patients, cases, predictions, explanations, reports, audit trails"],
            ["Storage", "Supabase Storage / local fallback", "Microscopy images and generated reports"],
        ],
    )

    add_heading(doc, "14. Diagram Atlas and Technical Interpretation", 1)
    add_paragraphs(
        doc,
        """The following figures are not decorative. Each one corresponds to a distinct validation or interpretability layer in the project. For a competition or thesis defense, these diagrams should be read as a map from architectural intent to measured evidence."""
    )
    add_figure(
        doc,
        PACKAGE_FIG_DIR / "Architecture_diagram.png",
        "Figure 1. Detailed system architecture diagram",
        """This figure is the top-to-bottom wiring view of PlasmaXAI. It shows the raw microscopy input flowing into dual backbones, morphology extraction, counterfactual construction, the learned fusion network, and then the explainability/reporting layer. The important conceptual reading is that counterfactual information sits inside the predictive graph rather than being bolted on after the fact.""",
        width=6.0,
    )
    add_figure(
        doc,
        FIG_DIR / "plasmaxai_fig13_framework_diagram.png",
        "Figure 2. PlasmaXAI framework overview",
        """This overview diagram compresses the research pipeline into major blocks. It is useful when explaining why the framework is novel: it contains learned fusion, counterfactual-conditioned image gating, patient-level analysis, and deployment translation in a single coherent system.""",
        width=6.0,
    )
    add_figure(
        doc,
        FIG_DIR / "novel_fig8_training_curves.png",
        "Figure 3. Training dynamics across loss, accuracy, recall, and AUC",
        """The left panel tracks optimization stability via train/validation loss. The middle panel shows whether the network generalizes rather than simply memorizing. The right panel is especially important because it tracks validation recall and validation AUC together, showing whether clinical sensitivity and ranking quality rise together during learning.""",
        width=6.1,
    )
    add_figure(
        doc,
        FIG_DIR / "plasmaxai_fig12_extended_model_comparison.png",
        "Figure 4. Extended model comparison against additional baselines",
        """This chart should be read as the framework-level justification for PlasmaXAI. It compares the final system not only against image backbones but also against morphology-only baselines. The important inference is that the final gains do not come from simply replacing one CNN with another; they emerge from structured multimodal fusion.""",
        width=6.1,
    )
    add_figure(
        doc,
        FIG_DIR / "plasmaxai_fig14_confusion_roc.png",
        "Figure 5. Confusion behavior and ROC geometry",
        """The confusion component makes the error tradeoff concrete, while the ROC panel shows global ranking behavior independent of a single threshold. Together, they explain why PlasmaXAI can preserve high AUC while also supporting threshold retuning for precision- or recall-oriented deployment.""",
        width=6.1,
    )
    add_figure(
        doc,
        FIG_DIR / "plasmaxai_fig4_counterfactual_explanation.png",
        "Figure 6. Counterfactual explanation diagram",
        """This figure exposes the internal evidence drivers. One panel shows learned modality weights, which reveals how much the sample relied on image versus morphology versus counterfactual signals. The second panel exposes the dominant counterfactual shifts, meaning which morphology coordinates would need to change most strongly to move the sample toward a benign-leaning state.""",
        width=6.1,
    )
    add_figure(
        doc,
        FIG_DIR / "plasmaxai_fig9_ablation_study.png",
        "Figure 7. Ablation study across image-only, no-morphology, no-counterfactual, and full system variants",
        """This is the most direct causal figure in the dossier. If one wants to know whether the counterfactual path genuinely matters, this figure answers that question. The key reading is that removing the counterfactual path harms plasma recall much more than removing morphology, which supports the design thesis that counterfactual conditioning helps difficult malignant-boundary cases.""",
        width=6.1,
    )
    add_figure(
        doc,
        FIG_DIR / "plasmaxai_fig10_calibration_bootstrap.png",
        "Figure 8. Calibration and bootstrap confidence intervals",
        """The calibration panel checks whether probability outputs behave like trustworthy confidence estimates rather than arbitrary scores. The bootstrap panel quantifies uncertainty around the headline metrics. For judge-facing or clinical-facing communication, this figure is what prevents the report from looking like it relies on a single lucky split.""",
        width=6.1,
    )
    add_figure(
        doc,
        FIG_DIR / "plasmaxai_fig6_clinical_insight.png",
        "Figure 9. Clinical insight and dominant patient-level shift signatures",
        """This figure translates cell-level mechanics into patient-level interpretation. The dominant shift features, especially the RGB-derived stain cues, show that the framework is detecting a coherent disease-associated morphology/stain program rather than random isolated decisions.""",
        width=6.1,
    )
    add_figure(
        doc,
        FIG_DIR / "novel_fig11_patient_signature_heatmap.png",
        "Figure 10. Patient signature heatmap",
        """The heatmap shows how top counterfactual features vary across patients. Read correctly, it is a between-patient consistency map: if diseased patients concentrate along similar feature bands, then the framework is capturing a stable cohort-level signature rather than independent noisy case guesses.""",
        width=6.1,
    )

    add_heading(doc, "15. Final Technical Interpretation", 1)
    add_paragraphs(
        doc,
        f"""Technically, PlasmaXAI is best understood as a five-block multimodal classifier in which counterfactual geometry controls part of the image-latent pathway itself. The reason the system is more than an ensemble is that it learns interactions between: (i) deep image embeddings, (ii) explicit morphology, (iii) counterfactual correction vectors, and (iv) branch-level probability summaries. That is a stricter and more integrated design than weighted averaging.

The results support that design. Compared with the best non-PlasmaXAI baseline in the current deployment table ({best_baseline['model']} at {best_baseline['accuracy']:.2f}% accuracy), the final framework achieves {plasmaxai_row['accuracy']:.2f}% accuracy, {plasmaxai_row['plasma_precision']:.2f}% plasma precision, {plasmaxai_row['plasma_recall']:.2f}% plasma recall, and {plasmaxai_row['auc']:.2f}% AUC. The bootstrap intervals stay tight, the ablation study penalizes the no-counterfactual model most strongly on plasma recall, and the patient-level signature analysis shows coherent disease-vs-normal separation. Taken together, those three layers of evidence - performance, ablation, and aggregation - are what make the framework technically defensible.

From an engineering standpoint, the same design also deploys cleanly. The model artifacts, scalers, and counterfactual bundle can be loaded by a FastAPI service, while the Next.js doctor workspace consumes structured outputs and exposes them as focus maps, interpretation text, patient timelines, editable reports, and downloadable PDFs. In other words, the project is not only a research pipeline; it is a research pipeline already translated into a product architecture."""
    )

    add_heading(doc, "16. Selected References and Artifact Anchors", 1)
    add_bullets(
        doc,
        [
            "Research script: research/scripts/novel_plasmaxai_pipeline.py",
            "Main summary artifact: research/outputs/novel/novel_summary.json",
            "Comparison artifact: research/outputs/novel/plasmaxai_extended_model_comparison.csv",
            "Ablation artifact: research/outputs/novel/plasmaxai_ablation_results.csv",
            "Patient summary artifact: research/outputs/novel/patient_counterfactual_summary.csv",
            "Inference service: apps/inference/app/predictor.py and apps/inference/app/main.py",
            "Web integration layer: apps/web/src/lib/inference/service.ts",
            "Database schema: apps/web/supabase/migrations/20260401_initial_schema.sql",
            "Architecture figure package: research/package/diagrams_and_images/paper_figures",
            "Project-level technical record: research/package/PlasmaXAI_ultra_detailed_full_record.txt",
        ],
    )

    return doc


def main() -> None:
    doc = build_document()
    OUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOC)
    print(OUT_DOC)


if __name__ == "__main__":
    main()
