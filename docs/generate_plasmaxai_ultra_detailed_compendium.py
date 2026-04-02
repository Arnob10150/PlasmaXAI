from __future__ import annotations

import json
import re
from pathlib import Path

import joblib
import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
RESEARCH_DIR = ROOT / "research"
PACKAGE_DIR = RESEARCH_DIR / "package"
NOVEL_DIR = RESEARCH_DIR / "outputs" / "novel"
FIG_DIR = NOVEL_DIR / "figures"
BASELINE_DIR = RESEARCH_DIR / "outputs" / "baseline"
OPT_DIR = RESEARCH_DIR / "outputs" / "optimization"
OUT_DOC = DOCS_DIR / "PlasmaXAI_Ultra_Detailed_Research_Architecture_Compendium.docx"
SOURCE_TXT = PACKAGE_DIR / "PlasmaXAI_ultra_detailed_full_record.txt"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    if "EquationBlock" not in doc.styles:
        style = doc.styles.add_style("EquationBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(10)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)

    if "CaptionSmall" not in doc.styles:
        style = doc.styles.add_style("CaptionSmall", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style.font.size = Pt(9)
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(8)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(11)


def add_paragraph(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    r.bold = bold


def add_equation(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="EquationBlock")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for index, value in enumerate(headers):
        hdr[index].text = value
        set_cell_shading(hdr[index], "1F2937")
        for paragraph in hdr[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def add_figure(doc: Document, path: Path, caption: str, note: str, width: float = 6.0) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(style="CaptionSmall")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.bold = True
    add_paragraph(doc, note)


def load_context() -> dict[str, object]:
    summary = json.loads((NOVEL_DIR / "novel_summary.json").read_text(encoding="utf-8"))
    bootstrap = json.loads((NOVEL_DIR / "bootstrap_summary.json").read_text(encoding="utf-8"))
    comparison_df = pd.read_csv(NOVEL_DIR / "plasmaxai_extended_model_comparison.csv")
    ablation_df = pd.read_csv(NOVEL_DIR / "plasmaxai_ablation_results.csv")
    patient_df = pd.read_csv(NOVEL_DIR / "patient_counterfactual_summary.csv")
    feature_bundle = joblib.load(NOVEL_DIR / "fusion_scalers.joblib")
    feature_blocks = feature_bundle["feature_blocks"]

    return {
        "summary": summary,
        "bootstrap": bootstrap,
        "comparison_df": comparison_df,
        "ablation_df": ablation_df,
        "patient_df": patient_df,
        "feature_dims": {
            "resnet": len(feature_blocks[0]),
            "densenet": len(feature_blocks[1]),
            "morphology": len(feature_blocks[2]),
            "counterfactual": len(feature_blocks[3]),
            "score": len(feature_blocks[4]),
        },
    }


def add_front_matter(doc: Document, ctx: dict[str, object]) -> None:
    summary = ctx["summary"]
    bootstrap = ctx["bootstrap"]
    comparison_df: pd.DataFrame = ctx["comparison_df"]
    ablation_df: pd.DataFrame = ctx["ablation_df"]
    dims = ctx["feature_dims"]

    plasmaxai_row = comparison_df.loc[comparison_df["model"] == "PlasmaXAI"].iloc[0]

    add_title(
        doc,
        "PlasmaXAI Ultra-Detailed Research, Architecture, and Equation Compendium",
        "Thesis-style technical record covering background, methodology, layer connectivity, equations, figures, and deployment translation",
    )

    add_paragraph(
        doc,
        "This document is designed as the deepest single-file technical reference for the PlasmaXAI project. It consolidates the long-form research record, the layer-by-layer novel architecture, the mathematics of the morphology and counterfactual branches, the benchmark and ablation tables, and a figure atlas tied to the saved research outputs.",
    )

    doc.add_heading("A. Executive Technical Snapshot", level=1)
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Primary problem", "Malignant plasma-cell recognition from microscopy images"],
            ["Final framework", "CounterfactualGuidedFusionNet / PlasmaXAI"],
            ["ResNet embedding dimension", str(dims["resnet"])],
            ["DenseNet embedding dimension", str(dims["densenet"])],
            ["Morphology feature dimension", str(dims["morphology"])],
            ["Counterfactual feature dimension", str(dims["counterfactual"])],
            ["Score block dimension", str(dims["score"])],
            ["Hidden dimension", str(summary["best_fusion_config"]["hidden_dim"])],
            ["Tabular dimension", str(summary["best_fusion_config"]["tabular_dim"])],
            ["Dropout", str(summary["best_fusion_config"]["dropout"])],
            ["Main novel run accuracy", f"{summary['novel_fusion_test_metrics']['accuracy']:.2f}%"],
            ["Main novel run plasma recall", f"{summary['novel_fusion_test_metrics']['plasma_recall']:.2f}%"],
            ["Deployment comparison accuracy", f"{plasmaxai_row['accuracy']:.2f}%"],
            ["Deployment comparison AUC", f"{plasmaxai_row['auc']:.2f}%"],
            ["Bootstrap accuracy mean", f"{bootstrap['novel']['accuracy']['mean']:.2f}%"],
            ["Bootstrap AUC mean", f"{bootstrap['novel']['auc']['mean']:.2f}%"],
        ],
    )

    doc.add_heading("B. Formal Mathematical Formulation", level=1)
    add_paragraph(
        doc,
        "The following equations formalize the major blocks of the PlasmaXAI research pipeline: morphology extraction, counterfactual boundary modeling, multimodal fusion, classifier scoring, and patient-level aggregation.",
    )

    equations = [
        "(1) x_s = [ p_resnet , p_densenet , p_cf ]",
        "(2) mu_R = (1 / |Omega|) * sum_{(u,v) in Omega} I_R(u,v)",
        "(3) mu_G = (1 / |Omega|) * sum_{(u,v) in Omega} I_G(u,v)",
        "(4) mu_B = (1 / |Omega|) * sum_{(u,v) in Omega} I_B(u,v)",
        "(5) I_stain = 255 - (1 / |Omega|) * sum_{(u,v) in Omega} Gray(u,v)",
        "(6) A_nuc = |Omega_nucleus|,   A_cell = |Omega_cell|,   A_cyto = max(A_cell - A_nuc, 0)",
        "(7) r_NC = A_nuc / (A_cyto + eps)",
        "(8) G_tex = std(Gray(Omega_cell))",
        "(9) R_shape = 4 * pi * A_cell / (P_cell^2 + eps)",
        "(10) e_r = f_resnet(I),   e_d = f_densenet(I)",
        "(11) p_resnet = sigma(w_r^T e_r + b_r),   p_densenet = sigma(w_d^T e_d + b_d)",
        "(12) z = StandardScaler(x_m)",
        "(13) m = w^T z + b,   p_cf = sigma(m)",
        "(14) delta_boundary = - ( max(m, 0) / ( ||w||_2^2 + eps ) ) * w",
        "(15) delta_proto = mu_benign - z",
        "(16) d_cf = max(m, 0) / ( ||w||_2 + eps )",
        "(17) r_hat = Dropout( GELU( W_r * LN(e_r) + b_r ) )",
        "(18) d_hat = Dropout( GELU( W_d * LN(e_d) + b_d ) )",
        "(19) m_hat = Dropout( GELU( W_m * LN(x_m) + b_m ) )",
        "(20) c_hat = Dropout( GELU( W_c * LN(x_cf) + b_c ) )",
        "(21) s_hat = Dropout( GELU( W_s * LN(x_s) + b_s ) )",
        "(22) g_r = sigmoid( W_cr * c_hat + b_cr ),   g_d = sigmoid( W_cd * c_hat + b_cd )",
        "(23) r_tilde = r_hat odot g_r,   d_tilde = d_hat odot g_d",
        "(24) s_tilde = Dropout( GELU( W_ms [m_hat ; c_hat] + b_ms ) ) + s_hat",
        "(25) alpha = softmax( W_gate * LN([r_tilde ; d_tilde ; m_hat ; c_hat ; s_tilde]) + b_gate )",
        "(26) u_modal = alpha_1 U_r(r_tilde) + alpha_2 U_d(d_tilde) + alpha_3 U_m(m_hat) + alpha_4 U_c(c_hat)",
        "(27) u = [ u_modal ; U_s(s_tilde) ]",
        "(28) y = W_o * Dropout( GELU( W_h * LN(u) + b_h ) ) + b_o",
        "(29) p = softmax(y)",
        "(30) L_CE = - sum_{i=1}^{N} sum_{k in {0,1}} w_k * 1[y_i = k] * log p_{ik}",
        "(31) PatientMeanProb_j = (1 / n_j) * sum_{i=1}^{n_j} p_i",
        "(32) PatientMeanCF_j = (1 / n_j^+) * sum_{i : plasma_i = 1} d_cf,i",
        "(33) Consistency_j = mean( cosine(Delta_i, Delta_k) ) for i != k",
        "(34) Score_patient = 0.6 * PatientMeanProb + 0.4 * PatientMeanCF",
    ]
    for equation in equations:
        add_equation(doc, equation)

    doc.add_heading("C. Layer-by-Layer Novel Architecture Summary", level=1)
    add_paragraph(
        doc,
        "The table below summarizes the concrete layer stack saved in the final CounterfactualGuidedFusionNet configuration. This is based on the actual research script and saved feature dimensions rather than generic assumptions.",
    )
    add_table(
        doc,
        ["Module", "Input dimension", "Layers", "Output dimension", "Functional role"],
        [
            ["Res encoder", str(dims["resnet"]), "LayerNorm -> Linear -> GELU -> Dropout", "256", "Compresses ResNet semantic embedding"],
            ["Dense encoder", str(dims["densenet"]), "LayerNorm -> Linear -> GELU -> Dropout", "256", "Compresses DenseNet texture embedding"],
            ["Morph encoder", str(dims["morphology"]), "LayerNorm -> Linear -> GELU -> Dropout", "128", "Encodes handcrafted morphology vector"],
            ["CF encoder", str(dims["counterfactual"]), "LayerNorm -> Linear -> GELU -> Dropout", "128", "Encodes counterfactual geometry"],
            ["Score encoder", str(dims["score"]), "LayerNorm -> Linear -> GELU -> Dropout", "128", "Encodes branch-level probabilities"],
            ["CF-to-Res gate", "128", "Linear -> Sigmoid", "256", "Gates ResNet latent with counterfactual evidence"],
            ["CF-to-Dense gate", "128", "Linear -> Sigmoid", "256", "Gates DenseNet latent with counterfactual evidence"],
            ["Morph-to-score refiner", "256", "Linear -> GELU -> Dropout", "128", "Injects morphology + CF context into score latent"],
            ["Modality gate", "896", "LayerNorm -> Linear -> Softmax", "4", "Learns relative modality weights"],
            ["Unification block", "256/128", "Linear projections to shared hidden space", "256", "Brings branches into common fusion space"],
            ["Classifier", "512", "LayerNorm -> Linear -> GELU -> Dropout -> Linear", "2", "Produces final binary logits"],
        ],
    )

    doc.add_heading("D. Quantitative Tables", level=1)
    comparison_rows = []
    comparison_df: pd.DataFrame = ctx["comparison_df"]
    for _, row in comparison_df.iterrows():
        comparison_rows.append(
            [
                str(row["model"]),
                f"{row['accuracy']:.2f}",
                f"{row['weighted_f1']:.2f}",
                f"{row['plasma_precision']:.2f}",
                f"{row['plasma_recall']:.2f}",
                f"{row['auc']:.2f}",
            ]
        )
    add_table(
        doc,
        ["Model", "Accuracy", "Weighted F1", "Plasma precision", "Plasma recall", "AUC"],
        comparison_rows,
    )
    add_paragraph(doc, "Ablation results are reproduced below because they provide the strongest evidence that the counterfactual path is functionally important rather than cosmetic.")
    ablation_rows = []
    for _, row in ablation_df.iterrows():
        ablation_rows.append(
            [
                str(row["label"]),
                f"{row['test_accuracy']:.2f}",
                f"{row['test_weighted_f1']:.2f}",
                f"{row['test_plasma_precision']:.2f}",
                f"{row['test_plasma_recall']:.2f}",
                f"{row['test_auc']:.2f}",
            ]
        )
    add_table(
        doc,
        ["Ablation", "Accuracy", "Weighted F1", "Plasma precision", "Plasma recall", "AUC"],
        ablation_rows,
    )


def add_figure_atlas(doc: Document) -> None:
    doc.add_heading("E. Figure and Diagram Atlas", level=1)
    add_paragraph(
        doc,
        "The images below are pulled from the actual saved research outputs. Each caption explains the role of the figure in the PlasmaXAI narrative, not just its filename.",
    )

    figures = [
        (
            FIG_DIR / "plasmaxai_fig13_framework_diagram.png",
            "Framework Diagram",
            "This is the high-level architecture view of PlasmaXAI. It is the best single figure for showing how microscopy input, dual deep backbones, morphology features, counterfactual features, and the learned fusion/classification path are connected end to end.",
        ),
        (
            FIG_DIR / "novel_fig8_training_curves.png",
            "Training Curves",
            "These curves show the optimization behavior of the novel fusion framework, including the evolution of loss, accuracy, and clinically important recall-oriented behavior across epochs.",
        ),
        (
            FIG_DIR / "plasmaxai_fig14_confusion_roc.png",
            "Confusion and ROC",
            "This figure ties threshold-level performance to threshold-independent ranking quality. The confusion panel grounds the error pattern, and the ROC panel explains why the framework remains separable across operating points.",
        ),
        (
            FIG_DIR / "plasmaxai_fig9_ablation_study.png",
            "Ablation Study",
            "This is one of the most important research figures because it demonstrates that removing the counterfactual path or morphology branch changes the performance profile. It supports the novelty claim experimentally.",
        ),
        (
            FIG_DIR / "novel_fig4_gates_counterfactuals.png",
            "Gate and Counterfactual Analysis",
            "This figure visualizes how the modality gates and counterfactual reasoning interact. It helps explain why the model is not just averaging branches blindly.",
        ),
        (
            FIG_DIR / "novel_fig11_patient_signature_heatmap.png",
            "Patient Signature Heatmap",
            "This figure summarizes patient-level consistency and cluster structure from the saved patient-level analysis. It moves the work beyond isolated cell classification.",
        ),
        (
            BASELINE_DIR / "fig6_counterfactual_analysis.png",
            "Notebook-Stage Counterfactual Analysis",
            "This is historically important because it shows the early explainability direction before counterfactual reasoning became part of the internal model path.",
        ),
        (
            OPT_DIR / "optimized_fig1_framework_overview.png",
            "Optimization-Stage Framework Overview",
            "This figure documents the stage between the notebook baseline and the final novel architecture. It helps show how the project matured methodologically.",
        ),
    ]

    for path, caption, note in figures:
        add_figure(doc, path, caption, note)


def add_full_record(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("F. Full Ultra-Detailed Research Record", level=1)
    add_paragraph(
        doc,
        "For completeness, the complete long-form research record is reproduced below in DOCX form. This section preserves the full narrative of the project, including motivation, architecture evolution, results, packaging, deployment notes, and limitations.",
    )

    lines = SOURCE_TXT.read_text(encoding="utf-8").splitlines()
    section_pattern = re.compile(r"^(\d+(?:\.\d+)*\.?)\s+(.+)$")

    paragraph_buffer: list[str] = []
    bullet_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_paragraph(doc, " ".join(line.strip() for line in paragraph_buffer if line.strip()))
            paragraph_buffer = []

    def flush_bullets() -> None:
        nonlocal bullet_buffer
        if bullet_buffer:
            add_bullets(doc, bullet_buffer)
            bullet_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()

        if not line or re.match(r"^-{5,}$", line) or re.match(r"^={5,}$", line):
            flush_paragraph()
            flush_bullets()
            continue

        if line.startswith("Purpose of this file") or line.startswith("Document intent"):
            flush_paragraph()
            flush_bullets()
            doc.add_heading(line, level=2)
            continue

        match = section_pattern.match(line)
        if match:
            flush_paragraph()
            flush_bullets()
            token = match.group(1).strip(".")
            depth = token.count(".") + 1
            level = 2 if depth == 1 else 3
            doc.add_heading(f"{token} {match.group(2)}", level=level)
            continue

        if line.startswith("- "):
            flush_paragraph()
            bullet_buffer.append(line[2:].strip())
            continue

        flush_bullets()
        paragraph_buffer.append(line)

    flush_paragraph()
    flush_bullets()


def build_document() -> Document:
    ctx = load_context()
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "PlasmaXAI Ultra-Detailed Research Architecture Compendium"
    doc.core_properties.author = "Arnob Aich Anurag / OpenAI Codex"

    add_front_matter(doc, ctx)
    add_figure_atlas(doc)
    add_full_record(doc)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUT_DOC)
    print(f"[ok] wrote {OUT_DOC}")


if __name__ == "__main__":
    main()
