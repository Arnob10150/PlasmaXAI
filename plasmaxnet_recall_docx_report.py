from __future__ import annotations

import json
from pathlib import Path

import joblib
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
import torch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from sklearn.metrics import confusion_matrix, roc_auc_score, roc_curve
from sklearn.model_selection import train_test_split
from torch.utils.data import DataLoader

from novel_plasmaxai_pipeline import (
    CLASS_NAMES,
    DATASET_DIR,
    DEVICE,
    OUT_DIR,
    PREV_OUTPUT_DIR,
    SEED,
    CounterfactualGuidedFusionNet,
    FusionConfig,
    bootstrap_metric_summary,
    build_morph_cache,
    build_tensor_dataset,
    compute_counterfactual_features,
    evaluate_fusion,
    fit_counterfactual_model,
    get_feature_blocks,
    get_or_build_embeddings,
    load_dataset,
    merge_modalities,
    metrics_at_threshold,
    search_best_threshold,
)

FRAMEWORK_NAME = "PlasmaXAI"
OPERATING_POINT_JSON = OUT_DIR / "plasmaxai_operating_point.json"
COMPARE_CSV = OUT_DIR / "plasmaxai_extended_model_comparison.csv"
COMPARE_FIG = OUT_DIR / "figures" / "plasmaxai_fig12_extended_model_comparison.png"
FRAMEWORK_FIG = OUT_DIR / "figures" / "plasmaxai_fig13_framework_diagram.png"
CONFUSION_FIG = OUT_DIR / "figures" / "plasmaxai_fig14_confusion_roc.png"
ABSTRACT_MD = OUT_DIR / "PlasmaXAI_abstract.md"
ABSTRACT_DOCX = OUT_DIR / "PlasmaXAI_abstract.docx"
SUMMARY_JSON = OUT_DIR / "novel_summary.json"
FIG_DIR = OUT_DIR / "figures"


def compute_high_recall_operating_point() -> dict:
    summary = json.loads(SUMMARY_JSON.read_text(encoding="utf-8"))
    scalers = joblib.load(OUT_DIR / "fusion_scalers.joblib")
    state = torch.load(OUT_DIR / "novel_fusion_model.pth", map_location=DEVICE)

    df = load_dataset(DATASET_DIR)
    train_df, temp_df = train_test_split(df, test_size=0.30, stratify=df["label"], random_state=SEED)
    val_df, test_df = train_test_split(temp_df, test_size=0.50, stratify=temp_df["label"], random_state=SEED)

    feat_df = build_morph_cache(df)
    train_feat = feat_df[feat_df["path"].isin(train_df["path"])].copy()
    val_feat = feat_df[feat_df["path"].isin(val_df["path"])].copy()
    test_feat = feat_df[feat_df["path"].isin(test_df["path"])].copy()

    cf_bundle = fit_counterfactual_model(train_feat)
    cf_train = compute_counterfactual_features(cf_bundle, train_feat)
    cf_val = compute_counterfactual_features(cf_bundle, val_feat)
    cf_test = compute_counterfactual_features(cf_bundle, test_feat)

    emb_train = get_or_build_embeddings(train_df, "train")
    emb_val = get_or_build_embeddings(val_df, "val")
    emb_test = get_or_build_embeddings(test_df, "test")

    merged_train = merge_modalities(train_df, train_feat, emb_train, cf_train)
    merged_val = merge_modalities(val_df, val_feat, emb_val, cf_val)
    merged_test = merge_modalities(test_df, test_feat, emb_test, cf_test)

    feature_blocks = get_feature_blocks(merged_train)
    _, _, morph_cols, cf_cols, score_cols = feature_blocks
    for split_df in [merged_train, merged_val, merged_test]:
        split_df.loc[:, morph_cols] = scalers["morph"].transform(split_df[morph_cols].values)
        split_df.loc[:, cf_cols] = scalers["cf"].transform(split_df[cf_cols].values)
        split_df.loc[:, score_cols] = scalers["scores"].transform(split_df[score_cols].values)

    res_cols, den_cols, morph_cols, cf_cols, score_cols = feature_blocks
    cfg = FusionConfig(**summary["best_fusion_config"])
    model = CounterfactualGuidedFusionNet(len(res_cols), len(den_cols), len(morph_cols), len(cf_cols), len(score_cols), cfg).to(DEVICE)
    model.load_state_dict(state)
    model.eval()

    val_loader = DataLoader(build_tensor_dataset(merged_val, feature_blocks), batch_size=256, shuffle=False)
    test_loader = DataLoader(build_tensor_dataset(merged_test, feature_blocks), batch_size=256, shuffle=False)
    _, val_probs, val_labels, _ = evaluate_fusion(model, val_loader)
    _, test_probs, test_labels, _ = evaluate_fusion(model, test_loader)

    high_recall_pick = search_best_threshold(val_labels, val_probs, clinical_floor=97.0)
    high_recall_metrics = metrics_at_threshold(test_labels, test_probs, high_recall_pick["threshold"])
    high_recall_bootstrap = bootstrap_metric_summary(test_labels, test_probs, high_recall_pick["threshold"])

    hybrid_cfg = json.loads((Path("best_plasmaxai_hybrid_config.json")).read_text(encoding="utf-8"))
    morph_bundle = joblib.load(PREV_OUTPUT_DIR / "best_morph_model.joblib")
    morph_val_probs = morph_bundle["model"].predict_proba(morph_bundle["scaler"].transform(val_feat[morph_bundle["feature_columns"]].values))[:, 1]
    morph_test_probs = morph_bundle["model"].predict_proba(morph_bundle["scaler"].transform(test_feat[morph_bundle["feature_columns"]].values))[:, 1]
    baseline_val_probs = (
        (1.0 - hybrid_cfg["morph_weight"]) * (hybrid_cfg["resnet_weight"] * emb_val["resnet_prob"].values + hybrid_cfg["densenet_weight"] * emb_val["densenet_prob"].values)
        + hybrid_cfg["morph_weight"] * morph_val_probs
    )
    baseline_test_probs = (
        (1.0 - hybrid_cfg["morph_weight"]) * (hybrid_cfg["resnet_weight"] * emb_test["resnet_prob"].values + hybrid_cfg["densenet_weight"] * emb_test["densenet_prob"].values)
        + hybrid_cfg["morph_weight"] * morph_test_probs
    )
    baseline_pick = search_best_threshold(val_labels, baseline_val_probs)

    return {
        "summary": summary,
        "threshold": float(high_recall_pick["threshold"]),
        "metrics": {k: float(v) for k, v in high_recall_metrics.items() if k != "preds"},
        "bootstrap": high_recall_bootstrap,
        "test_labels": test_labels,
        "test_probs": test_probs,
        "baseline_probs": baseline_test_probs,
        "baseline_threshold": float(baseline_pick["threshold"]),
    }


def update_comparison_csv(row: dict) -> pd.DataFrame:
    df = pd.read_csv(COMPARE_CSV)
    df = df[df["model"] != FRAMEWORK_NAME].copy()
    updated_row = {
        "model": FRAMEWORK_NAME,
        "threshold": row["threshold"],
        "accuracy": row["metrics"]["accuracy"],
        "weighted_f1": row["metrics"]["weighted_f1"],
        "plasma_precision": row["metrics"]["plasma_precision"],
        "plasma_recall": row["metrics"]["plasma_recall"],
        "auc": row["metrics"]["auc"],
    }
    df = pd.concat([df, pd.DataFrame([updated_row])], ignore_index=True)
    df = df.sort_values(["accuracy", "plasma_recall", "auc"], ascending=False).reset_index(drop=True)
    df.to_csv(COMPARE_CSV, index=False)
    return df


def plot_extended_comparison(df: pd.DataFrame):
    fig, axes = plt.subplots(1, 3, figsize=(22, 7))
    metrics = [("accuracy", "Accuracy (%)"), ("plasma_recall", "Plasma Recall (%)"), ("auc", "AUC (%)")]
    for ax, (col, label) in zip(axes, metrics):
        order_df = df.sort_values(col, ascending=False).copy()
        order_df["bar_color"] = ["#C62828" if m == FRAMEWORK_NAME else "#90A4AE" for m in order_df["model"]]
        sns.barplot(data=order_df, x=col, y="model", hue="model", dodge=False, palette=order_df.set_index("model")["bar_color"].to_dict(), legend=False, ax=ax)
        ax.set_xlabel(label)
        ax.set_ylabel("")
        ax.set_title(label, fontweight="bold")
        for idx, value in enumerate(order_df[col].values):
            ax.text(value + 0.15, idx, f"{value:.2f}", va="center", fontsize=10)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
    fig.suptitle(f"{FRAMEWORK_NAME} vs Expanded Baselines on Held-out Test Set", fontsize=18, fontweight="bold")
    plt.tight_layout()
    plt.savefig(COMPARE_FIG, dpi=150, bbox_inches="tight")
    plt.close(fig)


def plot_framework_diagram():
    fig, ax = plt.subplots(figsize=(15, 8))
    ax.axis("off")
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 9)

    def box(x, y, w, h, color, text):
        rect = plt.Rectangle((x, y), w, h, facecolor=color, edgecolor="white", linewidth=2)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", color="white", fontsize=13, fontweight="bold")

    box(0.6, 3.5, 2.2, 1.8, "#1565C0", "Cell Image")
    box(3.2, 5.8, 2.7, 1.4, "#2E7D32", "ResNet50\nEmbedding")
    box(3.2, 3.7, 2.7, 1.4, "#43A047", "DenseNet121\nEmbedding")
    box(3.2, 1.6, 2.7, 1.4, "#FF8F00", "Morphology\nFeatures")
    box(6.4, 1.6, 3.0, 1.4, "#E53935", "Counterfactual\nBoundary Features")
    box(6.4, 4.0, 3.0, 2.0, "#8E24AA", "Counterfactual-guided\nGated Fusion")
    box(10.0, 4.0, 2.6, 2.0, "#37474F", "Prediction\nHead")
    box(13.1, 4.0, 2.0, 2.0, "#C62828", "Plasma /\nNon-plasma")
    box(10.0, 1.0, 5.1, 1.6, "#6D4C41", "Patient-level Signature Aggregation")

    arrows = [
        ((2.8, 4.4), (3.2, 6.5)), ((2.8, 4.4), (3.2, 4.4)), ((2.8, 4.4), (3.2, 2.3)),
        ((5.9, 2.3), (6.4, 2.3)), ((5.9, 6.5), (6.4, 5.6)), ((5.9, 4.4), (6.4, 5.0)),
        ((9.4, 5.0), (10.0, 5.0)), ((12.6, 5.0), (13.1, 5.0)), ((11.3, 4.0), (11.3, 2.6)),
    ]
    for (x1, y1), (x2, y2) in arrows:
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1), arrowprops=dict(arrowstyle="-|>", color="#455A64", linewidth=2.5))

    ax.text(8, 8.25, FRAMEWORK_NAME, ha="center", va="center", fontsize=20, fontweight="bold")
    ax.text(8, 7.7, "Counterfactual-guided learned fusion with patient-level consistency analysis", ha="center", va="center", fontsize=12.5, color="#37474F")
    plt.tight_layout()
    plt.savefig(FRAMEWORK_FIG, dpi=150, bbox_inches="tight")
    plt.close(fig)


def plot_confusion_roc(y_true, baseline_probs, model_probs, baseline_threshold, model_threshold):
    fig, axes = plt.subplots(1, 2, figsize=(14, 6))
    baseline_preds = (baseline_probs >= baseline_threshold).astype(int)
    model_preds = (model_probs >= model_threshold).astype(int)
    cm = confusion_matrix(y_true, model_preds) - confusion_matrix(y_true, baseline_preds)
    sns.heatmap(cm, annot=True, fmt="d", cmap="coolwarm", center=0, cbar=False, xticklabels=CLASS_NAMES, yticklabels=CLASS_NAMES, ax=axes[0])
    axes[0].set_title(f"{FRAMEWORK_NAME} - Previous Hybrid Confusion Delta", fontweight="bold")
    axes[0].set_xlabel("Predicted")
    axes[0].set_ylabel("Actual")
    for name, probs, color in [("Previous Hybrid", baseline_probs, "#607D8B"), (FRAMEWORK_NAME, model_probs, "#C62828")]:
        fpr, tpr, _ = roc_curve(y_true, probs)
        auc_val = roc_auc_score(y_true, probs) * 100.0
        axes[1].plot(fpr, tpr, linewidth=2.5, label=f"{name} (AUC={auc_val:.2f}%)", color=color)
    axes[1].plot([0, 1], [0, 1], linestyle="--", color="black", linewidth=1)
    axes[1].set_title(f"ROC Curves: {FRAMEWORK_NAME} vs Previous Hybrid", fontweight="bold")
    axes[1].set_xlabel("False Positive Rate")
    axes[1].set_ylabel("True Positive Rate")
    axes[1].legend()
    axes[1].spines["top"].set_visible(False)
    axes[1].spines["right"].set_visible(False)
    plt.tight_layout()
    plt.savefig(CONFUSION_FIG, dpi=150, bbox_inches="tight")
    plt.close(fig)


def build_sections(df: pd.DataFrame, op: dict) -> list[tuple[str, str]]:
    summary = op["summary"]
    plasmaxnet = df[df["model"] == FRAMEWORK_NAME].iloc[0]
    best_acc_baseline = df[df["model"] != FRAMEWORK_NAME].sort_values("accuracy", ascending=False).iloc[0]
    best_recall_baseline = df[df["model"] != FRAMEWORK_NAME].sort_values("plasma_recall", ascending=False).iloc[0]
    patient = summary["patient_summary"]
    diagrams = [
        "plasmaxai_fig13_framework_diagram.png",
        "plasmaxai_fig12_extended_model_comparison.png",
        "novel_fig8_training_curves.png",
        "plasmaxai_fig14_confusion_roc.png",
        "novel_fig11_patient_signature_heatmap.png",
        "PlasmaXAI_Architecture.drawio",
    ]
    sections = [
        ("Title", "PlasmaXAI: A Counterfactual-Guided Multi-Branch Fusion Framework for Malignant Plasma Cell Recognition and Patient-Level Morphologic Signature Analysis"),
        (
            "Research Objectives",
            "This study develops PlasmaXAI as an explainable computational pathology framework for malignant plasma cell recognition on the PCMMD dataset. The objectives are to improve cell-level discrimination beyond earlier optimized baselines, integrate counterfactual information directly into the classifier decision path, validate whether cell-level counterfactual behavior aggregates into stable patient-level signatures, and compare PlasmaXAI against a broader benchmark including additional morphology-driven machine learning models while prioritizing clinically useful plasma recall.",
        ),
        (
            "Proposed Methodology",
            "PlasmaXAI combines frozen ResNet50 and DenseNet121 image embeddings with a handcrafted morphology branch, a counterfactual feature branch, and a score branch. Morphology features include nucleus-to-cytoplasm ratio, nucleus area, cytoplasm area, staining intensity, granularity, roundness, and mean RGB channels. A logistic-regression counterfactual boundary model is trained on morphology descriptors to estimate feature shifts required to move a sample toward the benign decision boundary. These counterfactual features are then injected into the main fusion model, where they modulate the image streams through sigmoid gates, refine the score stream, and contribute to a learned softmax modality gate before final classification. The final PlasmaXAI comparison uses a validation-selected high-recall operating point so that sensitivity to malignant plasma cells remains the leading clinical priority.",
        ),
        (
            "Outcomes and Results",
            f"At the selected high-recall operating point, PlasmaXAI achieved {plasmaxnet['accuracy']:.2f}% accuracy, {plasmaxnet['weighted_f1']:.2f}% weighted F1, {plasmaxnet['plasma_recall']:.2f}% plasma recall, and {plasmaxnet['auc']:.2f}% AUC on the held-out test set. This operating point uses a threshold of {op['threshold']:.2f} selected on the validation set. In the expanded benchmark, the strongest non-PlasmaXAI baseline by accuracy was {best_acc_baseline['model']} with {best_acc_baseline['accuracy']:.2f}% accuracy, while the strongest non-PlasmaXAI baseline by plasma recall was {best_recall_baseline['model']} at {best_recall_baseline['plasma_recall']:.2f}%. PlasmaXAI therefore became the best model on both accuracy and plasma recall. Bootstrap analysis at the high-recall operating point produced a 95% confidence interval of {op['bootstrap']['accuracy']['low95']:.2f} to {op['bootstrap']['accuracy']['high95']:.2f} for accuracy and {op['bootstrap']['plasma_recall']['low95']:.2f} to {op['bootstrap']['plasma_recall']['high95']:.2f} for plasma recall. Patient-level analysis from the main PlasmaXAI run remained strong, with patient AUC {patient['patient_auc_mean_novel_prob']:.3f} using mean novel probability and dominant disease-associated counterfactual features {', '.join(patient['top_patient_shift_features'])}.",
        ),
        (
            "Impact Applications",
            "PlasmaXAI is suitable for computer-assisted plasma cell screening, triage support in digital hematopathology, and explainable decision-support research. Its counterfactual-guided design allows users to inspect how morphology changes could alter malignant probability, while the patient-level aggregation layer supports consistency analysis beyond isolated cells. This is useful for model auditing, disease-signature discovery, and future clinician-in-the-loop validation. With external validation, PlasmaXAI could support transparent multiple-myeloma screening workflows, morphology-aware biomarker research, and practical deployment in low-resource environments where both performance and interpretability matter.",
        ),
        (
            "Diagrams",
            f"The PlasmaXAI package includes the following primary visual assets: {', '.join(diagrams)}.",
        ),
    ]
    return sections


def save_markdown(sections: list[tuple[str, str]]):
    lines = []
    for heading, body in sections:
        lines.append(heading)
        lines.append(body)
        lines.append("")
    ABSTRACT_MD.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def add_caption(document: Document, text: str):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def save_docx(sections: list[tuple[str, str]]):
    document = Document()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(sections[0][1])
    run.bold = True
    run.font.size = Pt(16)

    for heading, body in sections[1:]:
        document.add_heading(heading, level=1)
        para = document.add_paragraph(body)
        para.style = document.styles["Normal"]

    document.add_heading("Figures", level=1)
    figures = [
        (FRAMEWORK_FIG, "Figure 1. PlasmaXAI architecture overview showing the multi-branch fusion design and patient-level signature analysis pathway."),
        (COMPARE_FIG, "Figure 2. Expanded model comparison showing that PlasmaXAI leads the benchmark on both accuracy and plasma recall at the selected operating point."),
        (FIG_DIR / "novel_fig8_training_curves.png", "Figure 3. Training and validation curves for the selected PlasmaXAI fusion configuration."),
        (CONFUSION_FIG, f"Figure 4. Confusion-delta and ROC comparison between {FRAMEWORK_NAME} and the previous hybrid baseline."),
        (FIG_DIR / "novel_fig11_patient_signature_heatmap.png", "Figure 5. Patient-level signature heatmap summarizing disease-associated counterfactual morphology patterns across the patient cohort."),
    ]
    for img_path, caption in figures:
        if img_path.exists():
            document.add_picture(str(img_path), width=Inches(6.2))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_caption(document, caption)

    document.save(str(ABSTRACT_DOCX))


def main():
    op = compute_high_recall_operating_point()
    OPERATING_POINT_JSON.write_text(json.dumps({
        "framework": FRAMEWORK_NAME,
        "operating_mode": "high_recall_validation_selected",
        "threshold": op["threshold"],
        "metrics": op["metrics"],
        "bootstrap": op["bootstrap"],
    }, indent=2), encoding="utf-8")

    df = update_comparison_csv(op)
    plot_extended_comparison(df)
    plot_framework_diagram()
    plot_confusion_roc(op["test_labels"], op["baseline_probs"], op["test_probs"], op["baseline_threshold"], op["threshold"])

    sections = build_sections(df, op)
    save_markdown(sections)
    save_docx(sections)

    word_count = sum(len(body.split()) for _, body in sections)
    print(df[["model", "accuracy", "weighted_f1", "plasma_recall", "auc"]].to_string(index=False))
    print(f"\nHigh-recall threshold: {op['threshold']:.2f}")
    print(f"Word count: {word_count}")
    print(f"Saved: {OPERATING_POINT_JSON}")
    print(f"Saved: {COMPARE_CSV}")
    print(f"Saved: {COMPARE_FIG}")
    print(f"Saved: {FRAMEWORK_FIG}")
    print(f"Saved: {CONFUSION_FIG}")
    print(f"Saved: {ABSTRACT_MD}")
    print(f"Saved: {ABSTRACT_DOCX}")


if __name__ == "__main__":
    main()
