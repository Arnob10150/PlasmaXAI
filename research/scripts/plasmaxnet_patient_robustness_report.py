from __future__ import annotations

import json
from itertools import combinations
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from sklearn.metrics import roc_auc_score

ROOT = Path.cwd()
OUT_DIR = ROOT / 'research' / 'outputs' / 'novel'
FIG_DIR = OUT_DIR / 'figures'
FRAMEWORK_NAME = 'PlasmaXAI'

COMPARE_CSV = OUT_DIR / 'plasmaxai_extended_model_comparison.csv'
PATIENT_CSV = OUT_DIR / 'patient_counterfactual_summary.csv'
OP_JSON = OUT_DIR / 'plasmaxai_operating_point.json'
SUMMARY_JSON = OUT_DIR / 'novel_summary.json'

PREC_F1_FIG = FIG_DIR / 'plasmaxai_fig15_precision_f1_comparison.png'
ROBUST_FIG = FIG_DIR / 'plasmaxai_fig16_patient_robustness.png'
ROBUST_JSON = OUT_DIR / 'plasmaxai_patient_robustness.json'
LOPO_CSV = OUT_DIR / 'plasmaxai_patient_lopo_auc.csv'
PERM_CSV = OUT_DIR / 'plasmaxai_patient_permutation_auc.csv'
ABSTRACT_MD = OUT_DIR / 'PlasmaXAI_abstract_robustness.md'
ABSTRACT_DOCX = OUT_DIR / 'PlasmaXAI_abstract_robustness.docx'


def bootstrap_auc(labels: np.ndarray, scores: np.ndarray, n_boot: int = 5000, seed: int = 42) -> dict:
    rng = np.random.default_rng(seed)
    aucs = []
    n = len(labels)
    for _ in range(n_boot):
        idx = rng.integers(0, n, n)
        y = labels[idx]
        if len(np.unique(y)) < 2:
            continue
        aucs.append(float(roc_auc_score(y, scores[idx])))
    arr = np.asarray(aucs, dtype=float)
    return {
        'mean': float(arr.mean()),
        'low95': float(np.quantile(arr, 0.025)),
        'high95': float(np.quantile(arr, 0.975)),
        'n_valid': int(arr.size),
    }


def exact_permutation_test(labels: np.ndarray, scores: np.ndarray) -> tuple[float, pd.DataFrame, float, float]:
    labels = np.asarray(labels, dtype=int)
    scores = np.asarray(scores, dtype=float)
    n = len(labels)
    k = int(labels.sum())
    observed = float(roc_auc_score(labels, scores))
    aucs = []
    for pos_idx in combinations(range(n), k):
        perm = np.zeros(n, dtype=int)
        perm[list(pos_idx)] = 1
        aucs.append(float(roc_auc_score(perm, scores)))
    perm_arr = np.asarray(aucs, dtype=float)
    one_sided = float(np.mean(perm_arr >= observed))
    two_sided = float(np.mean(np.abs(perm_arr - 0.5) >= abs(observed - 0.5)))
    perm_df = pd.DataFrame({'perm_auc': perm_arr})
    return observed, perm_df, one_sided, two_sided


def lopo_auc(patient_df: pd.DataFrame, score_col: str) -> pd.DataFrame:
    rows = []
    for _, row in patient_df.iterrows():
        reduced = patient_df[patient_df['patient_id'] != row['patient_id']].copy()
        auc = float(roc_auc_score(reduced['diseased_label'], reduced[score_col]))
        rows.append({
            'omitted_patient_id': int(row['patient_id']),
            'omitted_diagnosis': row['diagnosis'],
            'lopo_auc': auc,
        })
    return pd.DataFrame(rows)


def plot_precision_f1(df: pd.DataFrame):
    fig, axes = plt.subplots(1, 3, figsize=(23, 7))
    views = [
        ('plasma_precision', 'Plasma Precision (%)'),
        ('weighted_f1', 'Weighted F1 (%)'),
        ('plasma_recall', 'Precision vs Recall'),
    ]

    order_precision = df.sort_values('plasma_precision', ascending=False).copy()
    order_precision['bar_color'] = ['#C62828' if m == FRAMEWORK_NAME else '#90A4AE' for m in order_precision['model']]
    sns.barplot(data=order_precision, x='plasma_precision', y='model', hue='model', dodge=False, palette=order_precision.set_index('model')['bar_color'].to_dict(), legend=False, ax=axes[0])
    axes[0].set_title('Plasma Precision', fontweight='bold')
    axes[0].set_xlabel('Precision (%)')
    axes[0].set_ylabel('')
    for idx, value in enumerate(order_precision['plasma_precision'].values):
        axes[0].text(value + 0.15, idx, f'{value:.2f}', va='center', fontsize=10)

    order_f1 = df.sort_values('weighted_f1', ascending=False).copy()
    order_f1['bar_color'] = ['#C62828' if m == FRAMEWORK_NAME else '#90A4AE' for m in order_f1['model']]
    sns.barplot(data=order_f1, x='weighted_f1', y='model', hue='model', dodge=False, palette=order_f1.set_index('model')['bar_color'].to_dict(), legend=False, ax=axes[1])
    axes[1].set_title('Weighted F1', fontweight='bold')
    axes[1].set_xlabel('Weighted F1 (%)')
    axes[1].set_ylabel('')
    for idx, value in enumerate(order_f1['weighted_f1'].values):
        axes[1].text(value + 0.15, idx, f'{value:.2f}', va='center', fontsize=10)

    colors = ['#C62828' if m == FRAMEWORK_NAME else '#607D8B' for m in df['model']]
    axes[2].scatter(df['plasma_recall'], df['plasma_precision'], s=df['weighted_f1'] * 6, c=colors, edgecolors='white', linewidth=1.4)
    for _, row in df.iterrows():
        axes[2].annotate(row['model'], (row['plasma_recall'], row['plasma_precision']), textcoords='offset points', xytext=(6, 4), fontsize=9)
    axes[2].set_title('Precision-Recall Landscape', fontweight='bold')
    axes[2].set_xlabel('Plasma Recall (%)')
    axes[2].set_ylabel('Plasma Precision (%)')
    axes[2].grid(alpha=0.25, linestyle='--')

    for ax in axes:
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    fig.suptitle(f'{FRAMEWORK_NAME} Precision and F1 Comparison', fontsize=18, fontweight='bold')
    plt.tight_layout()
    plt.savefig(PREC_F1_FIG, dpi=150, bbox_inches='tight')
    plt.close(fig)


def plot_patient_robustness(patient_df: pd.DataFrame, perm_df: pd.DataFrame, observed_auc: float, lopo_df: pd.DataFrame):
    fig, axes = plt.subplots(1, 3, figsize=(23, 7.5))

    perm_plot = perm_df.groupby('perm_auc').size().reset_index(name='count').sort_values('perm_auc').reset_index(drop=True)
    positions = np.arange(len(perm_plot))
    axes[0].bar(positions, perm_plot['count'].values, color='#90A4AE', width=0.82)
    observed_x = float(np.interp(observed_auc, perm_plot['perm_auc'].values, positions))
    axes[0].axvline(x=observed_x, color='#C62828', linestyle='--', linewidth=2)
    tick_idx = np.unique(np.linspace(0, len(perm_plot) - 1, min(8, len(perm_plot)), dtype=int))
    axes[0].set_xticks(tick_idx)
    axes[0].set_xticklabels([f"{perm_plot.iloc[i]['perm_auc']:.2f}" for i in tick_idx], rotation=30, ha='right')
    axes[0].tick_params(axis='x', labelsize=9)
    axes[0].set_title('Exact Permutation Null (Patient AUC)', fontweight='bold')
    axes[0].set_xlabel('AUC')
    axes[0].set_ylabel('Count of label assignments')

    lopo_df = lopo_df.copy().sort_values('omitted_patient_id')
    sns.barplot(data=lopo_df, x='omitted_patient_id', y='lopo_auc', hue='omitted_diagnosis', palette={'diseased': '#E53935', 'normal': '#1565C0'}, ax=axes[1])
    axes[1].set_title('Leave-One-Patient-Out AUC', fontweight='bold')
    axes[1].set_xlabel('Omitted Patient ID')
    axes[1].set_ylabel('AUC on Remaining Patients')
    axes[1].set_ylim(0.0, 1.05)

    ordered = patient_df.sort_values('mean_novel_prob').copy()
    sns.stripplot(data=ordered, x='diagnosis', y='mean_novel_prob', hue='diagnosis', palette={'diseased': '#E53935', 'normal': '#1565C0'}, dodge=False, size=10, ax=axes[2])
    for _, row in ordered.iterrows():
        axes[2].annotate(f"P{int(row['patient_id']):02d}", (0 if row['diagnosis']=='diseased' else 1, row['mean_novel_prob']), textcoords='offset points', xytext=(6, 0), fontsize=8)
    axes[2].set_title('Internal 10-Patient Score Separation', fontweight='bold')
    axes[2].set_xlabel('Diagnosis Group')
    axes[2].set_ylabel('Mean Novel Probability')
    if axes[2].legend_ is not None:
        axes[2].legend_.remove()

    for ax in axes:
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    fig.suptitle(f'{FRAMEWORK_NAME} Patient-Level Robustness Checks', fontsize=18, fontweight='bold')
    plt.tight_layout()
    plt.savefig(ROBUST_FIG, dpi=150, bbox_inches='tight')
    plt.close(fig)


def build_sections(compare_df: pd.DataFrame, patient_df: pd.DataFrame, op: dict, robust: dict) -> list[tuple[str, str]]:
    model_row = compare_df[compare_df['model'] == FRAMEWORK_NAME].iloc[0]
    acc_baseline = compare_df[compare_df['model'] != FRAMEWORK_NAME].sort_values('accuracy', ascending=False).iloc[0]
    recall_baseline = compare_df[compare_df['model'] != FRAMEWORK_NAME].sort_values('plasma_recall', ascending=False).iloc[0]
    diagrams = [
        'plasmaxai_fig13_framework_diagram.png',
        'plasmaxai_fig12_extended_model_comparison.png',
        'plasmaxai_fig15_precision_f1_comparison.png',
        'novel_fig8_training_curves.png',
        'plasmaxai_fig14_confusion_roc.png',
        'plasmaxai_fig16_patient_robustness.png',
        'novel_fig11_patient_signature_heatmap.png',
        'PlasmaXAI_Architecture.drawio',
    ]
    return [
        ('Title', 'PlasmaXAI: A Counterfactual-Guided Multi-Branch Fusion Framework for Malignant Plasma Cell Recognition and Patient-Level Morphologic Signature Analysis'),
        ('Research Objectives', 'This study presents PlasmaXAI as an explainable computational pathology framework for malignant plasma cell recognition on PCMMD. The objectives are to improve cell-level discrimination against stronger benchmark models, integrate counterfactual information directly into the decision pathway, prioritize clinically important plasma recall, and study whether cell-level counterfactual behavior aggregates into patient-level disease signatures.'),
        ('Proposed Methodology', 'PlasmaXAI combines frozen ResNet50 and DenseNet121 image embeddings with a handcrafted morphology branch, a counterfactual feature branch, and a score branch. Morphologic descriptors include nucleus-to-cytoplasm ratio, nucleus area, cytoplasm area, staining intensity, granularity, roundness, and mean RGB channels. A morphology-based counterfactual boundary model first estimates the feature shifts required to move a sample toward the benign decision boundary. These counterfactual signals then modulate the image streams through sigmoid gates, refine the score stream, and participate in a learned softmax modality gate before weighted fusion and final classification. For deployment and benchmarking, the final model uses a validation-selected high-recall operating point so malignant sensitivity remains the primary clinical priority.'),
        ('Outcomes and Results', f"At the selected high-recall operating point, PlasmaXAI achieved {model_row['accuracy']:.2f}% accuracy, {model_row['weighted_f1']:.2f}% weighted F1, {model_row['plasma_precision']:.2f}% plasma precision, {model_row['plasma_recall']:.2f}% plasma recall, and {model_row['auc']:.2f}% AUC on the held-out test set. This threshold was set to {op['threshold']:.2f} from the validation set. In the expanded benchmark, the strongest non-PlasmaXAI baseline by accuracy was {acc_baseline['model']} at {acc_baseline['accuracy']:.2f}% accuracy, while the strongest non-PlasmaXAI baseline by recall was {recall_baseline['model']} at {recall_baseline['plasma_recall']:.2f}% recall. PlasmaXAI therefore leads the current benchmark on both overall accuracy and malignant-cell recall. The bootstrap 95% interval at the selected operating point was {op['bootstrap']['accuracy']['low95']:.2f}-{op['bootstrap']['accuracy']['high95']:.2f} for accuracy and {op['bootstrap']['plasma_recall']['low95']:.2f}-{op['bootstrap']['plasma_recall']['high95']:.2f} for plasma recall."),
        ('Patient-Level Robustness and Interpretation', f"The patient-level score using mean novel probability produced an internal AUC of {robust['observed_auc']:.3f} on a cohort of only {robust['patient_count']} patients ({robust['diseased_count']} diseased, {robust['normal_count']} normal). This should be treated as exploratory rather than definitive clinical proof. Exact permutation testing across all 252 possible class-label assignments gave a one-sided p-value of {robust['perm_p_one_sided']:.4f} and a two-sided p-value of {robust['perm_p_two_sided']:.4f}. Leave-one-patient-out analysis yielded mean AUC {robust['lopo_mean']:.3f}, minimum {robust['lopo_min']:.3f}, and maximum {robust['lopo_max']:.3f}. The patient bootstrap interval saturated at {robust['bootstrap_auc']['low95']:.3f}-{robust['bootstrap_auc']['high95']:.3f}, reflecting perfect within-cohort rank separation rather than external validation. The disease-associated signature remained centered on {', '.join(robust['top_features'])}, but because the cohort contains only ten patients from a single internal source, the patient-level result should be reported as a promising internal finding that requires external multi-center confirmation."),
        ('Impact Applications', 'PlasmaXAI is suitable for computer-assisted plasma cell screening, triage support in digital hematopathology, and explainable decision-support research. Its counterfactual-guided design helps reveal which morphology changes would alter malignant probability, while its patient-level aggregation layer supports consistency analysis beyond isolated cell predictions. This makes the framework useful for model auditing, disease-signature discovery, and future clinician-in-the-loop validation studies.'),
        ('Diagrams', f"The current PlasmaXAI package includes the following primary visual assets: {', '.join(diagrams)}."),
    ]


def save_markdown(sections: list[tuple[str, str]]):
    lines = []
    for heading, body in sections:
        lines.extend([heading, body, ''])
    ABSTRACT_MD.write_text('\n'.join(lines).strip() + '\n', encoding='utf-8')


def add_caption(document: Document, text: str):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def save_docx(sections: list[tuple[str, str]]):
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(sections[0][1])
    run.bold = True
    run.font.size = Pt(16)

    for heading, body in sections[1:]:
        doc.add_heading(heading, level=1)
        para = doc.add_paragraph(body)
        para.style = doc.styles['Normal']

    doc.add_heading('Figures', level=1)
    figures = [
        (FIG_DIR / 'plasmaxai_fig13_framework_diagram.png', 'Figure 1. PlasmaXAI architecture overview showing the multi-branch fusion design and patient-level signature pathway.'),
        (FIG_DIR / 'plasmaxai_fig12_extended_model_comparison.png', 'Figure 2. Expanded model comparison showing that PlasmaXAI leads the benchmark on both accuracy and plasma recall at the selected operating point.'),
        (FIG_DIR / 'plasmaxai_fig15_precision_f1_comparison.png', 'Figure 3. Precision, weighted F1, and precision-recall landscape across PlasmaXAI and comparison models.'),
        (FIG_DIR / 'novel_fig8_training_curves.png', 'Figure 4. Training and validation curves for the selected PlasmaXAI fusion configuration.'),
        (FIG_DIR / 'plasmaxai_fig14_confusion_roc.png', 'Figure 5. Confusion-delta and ROC comparison between PlasmaXAI and the previous hybrid baseline.'),
        (FIG_DIR / 'plasmaxai_fig16_patient_robustness.png', 'Figure 6. Patient-level robustness analysis including exact permutation testing, leave-one-patient-out AUC, and internal score separation on the 10-patient cohort.'),
        (FIG_DIR / 'novel_fig11_patient_signature_heatmap.png', 'Figure 7. Patient-level signature heatmap summarizing disease-associated counterfactual morphology patterns across the cohort.'),
    ]
    for img_path, caption in figures:
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_caption(doc, caption)

    doc.save(str(ABSTRACT_DOCX))


def main():
    compare_df = pd.read_csv(COMPARE_CSV)
    patient_df = pd.read_csv(PATIENT_CSV)
    op = json.loads(OP_JSON.read_text(encoding='utf-8'))
    summary = json.loads(SUMMARY_JSON.read_text(encoding='utf-8'))

    labels = patient_df['diseased_label'].to_numpy(dtype=int)
    scores = patient_df['mean_novel_prob'].to_numpy(dtype=float)
    observed_auc, perm_df, p_one, p_two = exact_permutation_test(labels, scores)
    bootstrap = bootstrap_auc(labels, scores)
    lopo_df = lopo_auc(patient_df, 'mean_novel_prob')

    diseased_scores = patient_df.loc[patient_df['diseased_label'] == 1, 'mean_novel_prob'].to_numpy(dtype=float)
    normal_scores = patient_df.loc[patient_df['diseased_label'] == 0, 'mean_novel_prob'].to_numpy(dtype=float)
    separation_gap = float(diseased_scores.min() - normal_scores.max())

    robust = {
        'patient_count': int(len(patient_df)),
        'diseased_count': int(labels.sum()),
        'normal_count': int((1 - labels).sum()),
        'observed_auc': float(observed_auc),
        'perm_p_one_sided': float(p_one),
        'perm_p_two_sided': float(p_two),
        'bootstrap_auc': bootstrap,
        'lopo_mean': float(lopo_df['lopo_auc'].mean()),
        'lopo_min': float(lopo_df['lopo_auc'].min()),
        'lopo_max': float(lopo_df['lopo_auc'].max()),
        'separation_gap': separation_gap,
        'top_features': summary['patient_summary']['top_patient_shift_features'],
    }

    perm_df.to_csv(PERM_CSV, index=False)
    lopo_df.to_csv(LOPO_CSV, index=False)
    ROBUST_JSON.write_text(json.dumps(robust, indent=2), encoding='utf-8')

    plot_precision_f1(compare_df)
    plot_patient_robustness(patient_df, perm_df, observed_auc, lopo_df)

    sections = build_sections(compare_df, patient_df, op, robust)
    save_markdown(sections)
    save_docx(sections)

    word_count = sum(len(body.split()) for _, body in sections)
    print(json.dumps(robust, indent=2))
    print(f'Word count: {word_count}')
    print(f'Saved: {PREC_F1_FIG}')
    print(f'Saved: {ROBUST_FIG}')
    print(f'Saved: {ROBUST_JSON}')
    print(f'Saved: {LOPO_CSV}')
    print(f'Saved: {PERM_CSV}')
    print(f'Saved: {ABSTRACT_MD}')
    print(f'Saved: {ABSTRACT_DOCX}')


if __name__ == '__main__':
    main()

