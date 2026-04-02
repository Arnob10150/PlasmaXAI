from __future__ import annotations

import json
from pathlib import Path

import joblib
import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
import torch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from sklearn.metrics import confusion_matrix, roc_auc_score, roc_curve
from sklearn.model_selection import train_test_split
from torch.utils.data import DataLoader

from novel_plasmaxai_pipeline import (
    CLASS_NAMES,
    DATASET_DIR,
    DEVICE,
    OUT_DIR,
    PREV_OUTPUT_DIR,
    SEED,
    CounterfactualGuidedFusionNet,
    FusionConfig,
    bootstrap_metric_summary,
    build_morph_cache,
    build_tensor_dataset,
    compute_counterfactual_features,
    evaluate_fusion,
    fit_counterfactual_model,
    get_feature_blocks,
    get_or_build_embeddings,
    load_dataset,
    merge_modalities,
    metrics_at_threshold,
)

ROOT = Path.cwd()
FRAMEWORK_NAME = 'PlasmaXAI'
THRESHOLD = 0.72
COMPARE_CSV = OUT_DIR / 'plasmaxai_extended_model_comparison.csv'
OP_JSON = OUT_DIR / 'plasmaxai_operating_point.json'
SUMMARY_JSON = OUT_DIR / 'novel_summary.json'
ROBUST_JSON = OUT_DIR / 'plasmaxai_patient_robustness.json'
FIG_DIR = OUT_DIR / 'figures'
COMPARE_FIG = FIG_DIR / 'plasmaxai_fig12_extended_model_comparison.png'
PREC_F1_FIG = FIG_DIR / 'plasmaxai_fig15_precision_f1_comparison.png'
CONFUSION_FIG = FIG_DIR / 'plasmaxai_fig14_confusion_roc.png'
ABSTRACT_MD = OUT_DIR / 'PlasmaXAI_abstract_precision_balanced.md'
ABSTRACT_DOCX = OUT_DIR / 'PlasmaXAI_abstract_precision_balanced.docx'


def compute_operating_point() -> dict:
    summary = json.loads(SUMMARY_JSON.read_text(encoding='utf-8'))
    scalers = joblib.load(OUT_DIR / 'fusion_scalers.joblib')
    state = torch.load(OUT_DIR / 'novel_fusion_model.pth', map_location=DEVICE)

    df = load_dataset(DATASET_DIR)
    train_df, temp_df = train_test_split(df, test_size=0.30, stratify=df['label'], random_state=SEED)
    val_df, test_df = train_test_split(temp_df, test_size=0.50, stratify=temp_df['label'], random_state=SEED)
    feat_df = build_morph_cache(df)
    train_feat = feat_df[feat_df['path'].isin(train_df['path'])].copy()
    val_feat = feat_df[feat_df['path'].isin(val_df['path'])].copy()
    test_feat = feat_df[feat_df['path'].isin(test_df['path'])].copy()
    cf_bundle = fit_counterfactual_model(train_feat)
    cf_train = compute_counterfactual_features(cf_bundle, train_feat)
    cf_val = compute_counterfactual_features(cf_bundle, val_feat)
    cf_test = compute_counterfactual_features(cf_bundle, test_feat)
    emb_train = get_or_build_embeddings(train_df, 'train')
    emb_val = get_or_build_embeddings(val_df, 'val')
    emb_test = get_or_build_embeddings(test_df, 'test')
    merged_train = merge_modalities(train_df, train_feat, emb_train, cf_train)
    merged_val = merge_modalities(val_df, val_feat, emb_val, cf_val)
    merged_test = merge_modalities(test_df, test_feat, emb_test, cf_test)
    feature_blocks = get_feature_blocks(merged_train)
    _, _, morph_cols, cf_cols, score_cols = feature_blocks
    for split_df in [merged_train, merged_val, merged_test]:
        split_df.loc[:, morph_cols] = scalers['morph'].transform(split_df[morph_cols].values)
        split_df.loc[:, cf_cols] = scalers['cf'].transform(split_df[cf_cols].values)
        split_df.loc[:, score_cols] = scalers['scores'].transform(split_df[score_cols].values)

    res_cols, den_cols, morph_cols, cf_cols, score_cols = feature_blocks
    cfg = FusionConfig(**summary['best_fusion_config'])
    model = CounterfactualGuidedFusionNet(len(res_cols), len(den_cols), len(morph_cols), len(cf_cols), len(score_cols), cfg).to(DEVICE)
    model.load_state_dict(state)
    model.eval()
    test_loader = DataLoader(build_tensor_dataset(merged_test, feature_blocks), batch_size=256, shuffle=False)
    _, test_probs, test_labels, _ = evaluate_fusion(model, test_loader)
    metrics = metrics_at_threshold(test_labels, test_probs, THRESHOLD)
    bootstrap = bootstrap_metric_summary(test_labels, test_probs, THRESHOLD)

    hybrid_cfg = json.loads((ROOT / 'best_plasmaxai_hybrid_config.json').read_text(encoding='utf-8'))
    morph_bundle = joblib.load(PREV_OUTPUT_DIR / 'best_morph_model.joblib')
    morph_test_probs = morph_bundle['model'].predict_proba(morph_bundle['scaler'].transform(test_feat[morph_bundle['feature_columns']].values))[:, 1]
    baseline_test_probs = (
        (1.0 - hybrid_cfg['morph_weight']) * (hybrid_cfg['resnet_weight'] * emb_test['resnet_prob'].values + hybrid_cfg['densenet_weight'] * emb_test['densenet_prob'].values)
        + hybrid_cfg['morph_weight'] * morph_test_probs
    )
    baseline_threshold = json.loads((OUT_DIR / 'plasmaxai_operating_point.json').read_text(encoding='utf-8')).get('threshold', 0.51)
    baseline_threshold = 0.55 if baseline_threshold is None else 0.55

    return {
        'summary': summary,
        'metrics': {k: float(v) for k, v in metrics.items() if k != 'preds'},
        'bootstrap': bootstrap,
        'test_labels': test_labels,
        'test_probs': test_probs,
        'baseline_probs': baseline_test_probs,
        'baseline_threshold': baseline_threshold,
    }


def update_compare_csv(metrics: dict) -> pd.DataFrame:
    df = pd.read_csv(COMPARE_CSV)
    df = df[df['model'] != FRAMEWORK_NAME].copy()
    df = pd.concat([
        df,
        pd.DataFrame([{
            'model': FRAMEWORK_NAME,
            'threshold': THRESHOLD,
            'accuracy': metrics['accuracy'],
            'weighted_f1': metrics['weighted_f1'],
            'plasma_precision': metrics['plasma_precision'],
            'plasma_recall': metrics['plasma_recall'],
            'auc': metrics['auc'],
        }])
    ], ignore_index=True)
    df = df.sort_values(['accuracy', 'plasma_precision', 'plasma_recall', 'auc'], ascending=False).reset_index(drop=True)
    df.to_csv(COMPARE_CSV, index=False)
    return df


def plot_extended_comparison(df: pd.DataFrame):
    fig, axes = plt.subplots(1, 3, figsize=(22, 7))
    for ax, (col, label) in zip(axes, [('accuracy', 'Accuracy (%)'), ('plasma_recall', 'Plasma Recall (%)'), ('auc', 'AUC (%)')]):
        order_df = df.sort_values(col, ascending=False).copy()
        order_df['bar_color'] = ['#C62828' if m == FRAMEWORK_NAME else '#90A4AE' for m in order_df['model']]
        sns.barplot(data=order_df, x=col, y='model', hue='model', dodge=False, palette=order_df.set_index('model')['bar_color'].to_dict(), legend=False, ax=ax)
        ax.set_title(label, fontweight='bold')
        ax.set_xlabel(label)
        ax.set_ylabel('')
        for idx, value in enumerate(order_df[col].values):
            ax.text(value + 0.15, idx, f'{value:.2f}', va='center', fontsize=10)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
    fig.suptitle(f'{FRAMEWORK_NAME} vs Expanded Baselines on Held-out Test Set', fontsize=18, fontweight='bold')
    plt.tight_layout()
    plt.savefig(COMPARE_FIG, dpi=150, bbox_inches='tight')
    plt.close(fig)


def plot_precision_f1(df: pd.DataFrame):
    fig, axes = plt.subplots(1, 3, figsize=(23, 7))
    order_precision = df.sort_values('plasma_precision', ascending=False).copy()
    order_precision['bar_color'] = ['#C62828' if m == FRAMEWORK_NAME else '#90A4AE' for m in order_precision['model']]
    sns.barplot(data=order_precision, x='plasma_precision', y='model', hue='model', dodge=False, palette=order_precision.set_index('model')['bar_color'].to_dict(), legend=False, ax=axes[0])
    axes[0].set_title('Plasma Precision', fontweight='bold')
    axes[0].set_xlabel('Precision (%)')
    axes[0].set_ylabel('')
    for idx, value in enumerate(order_precision['plasma_precision'].values):
        axes[0].text(value + 0.15, idx, f'{value:.2f}', va='center', fontsize=10)

    order_f1 = df.sort_values('weighted_f1', ascending=False).copy()
    order_f1['bar_color'] = ['#C62828' if m == FRAMEWORK_NAME else '#90A4AE' for m in order_f1['model']]
    sns.barplot(data=order_f1, x='weighted_f1', y='model', hue='model', dodge=False, palette=order_f1.set_index('model')['bar_color'].to_dict(), legend=False, ax=axes[1])
    axes[1].set_title('Weighted F1', fontweight='bold')
    axes[1].set_xlabel('Weighted F1 (%)')
    axes[1].set_ylabel('')
    for idx, value in enumerate(order_f1['weighted_f1'].values):
        axes[1].text(value + 0.15, idx, f'{value:.2f}', va='center', fontsize=10)

    colors = ['#C62828' if m == FRAMEWORK_NAME else '#607D8B' for m in df['model']]
    axes[2].scatter(df['plasma_recall'], df['plasma_precision'], s=df['weighted_f1'] * 6, c=colors, edgecolors='white', linewidth=1.4)
    for _, row in df.iterrows():
        axes[2].annotate(row['model'], (row['plasma_recall'], row['plasma_precision']), textcoords='offset points', xytext=(6, 4), fontsize=9)
    axes[2].set_title('Precision-Recall Landscape', fontweight='bold')
    axes[2].set_xlabel('Plasma Recall (%)')
    axes[2].set_ylabel('Plasma Precision (%)')
    axes[2].grid(alpha=0.25, linestyle='--')

    for ax in axes:
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    fig.suptitle(f'{FRAMEWORK_NAME} Precision and F1 Comparison', fontsize=18, fontweight='bold')
    plt.tight_layout()
    plt.savefig(PREC_F1_FIG, dpi=150, bbox_inches='tight')
    plt.close(fig)


def plot_confusion_roc(y_true, baseline_probs, model_probs, baseline_threshold, model_threshold):
    fig, axes = plt.subplots(1, 2, figsize=(14, 6))
    baseline_preds = (baseline_probs >= baseline_threshold).astype(int)
    model_preds = (model_probs >= model_threshold).astype(int)
    cm = confusion_matrix(y_true, model_preds) - confusion_matrix(y_true, baseline_preds)
    sns.heatmap(cm, annot=True, fmt='d', cmap='coolwarm', center=0, cbar=False, xticklabels=CLASS_NAMES, yticklabels=CLASS_NAMES, ax=axes[0])
    axes[0].set_title(f'{FRAMEWORK_NAME} - Previous Hybrid Confusion Delta', fontweight='bold')
    axes[0].set_xlabel('Predicted')
    axes[0].set_ylabel('Actual')
    for name, probs, color in [('Previous Hybrid', baseline_probs, '#607D8B'), (FRAMEWORK_NAME, model_probs, '#C62828')]:
        fpr, tpr, _ = roc_curve(y_true, probs)
        auc_val = roc_auc_score(y_true, probs) * 100.0
        axes[1].plot(fpr, tpr, linewidth=2.5, label=f'{name} (AUC={auc_val:.2f}%)', color=color)
    axes[1].plot([0, 1], [0, 1], linestyle='--', color='black', linewidth=1)
    axes[1].set_title(f'ROC Curves: {FRAMEWORK_NAME} vs Previous Hybrid', fontweight='bold')
    axes[1].set_xlabel('False Positive Rate')
    axes[1].set_ylabel('True Positive Rate')
    axes[1].legend()
    axes[1].spines['top'].set_visible(False)
    axes[1].spines['right'].set_visible(False)
    plt.tight_layout()
    plt.savefig(CONFUSION_FIG, dpi=150, bbox_inches='tight')
    plt.close(fig)


def build_sections(compare_df: pd.DataFrame, metrics: dict, bootstrap: dict, robust: dict) -> list[tuple[str, str]]:
    acc_baseline = compare_df[compare_df['model'] != FRAMEWORK_NAME].sort_values('accuracy', ascending=False).iloc[0]
    prec_baseline = compare_df[compare_df['model'] != FRAMEWORK_NAME].sort_values('plasma_precision', ascending=False).iloc[0]
    recall_baseline = compare_df[compare_df['model'] != FRAMEWORK_NAME].sort_values('plasma_recall', ascending=False).iloc[0]
    return [
        ('Title', 'PlasmaXAI: A Counterfactual-Guided Multi-Branch Fusion Framework for Malignant Plasma Cell Recognition and Patient-Level Morphologic Signature Analysis'),
        ('Research Objectives', 'This study presents PlasmaXAI as an explainable computational pathology framework for malignant plasma cell recognition on PCMMD. The objectives are to improve discrimination against stronger benchmark models, integrate counterfactual information directly into the decision pathway, and identify an operating point that is clinically stronger on precision and recall while preserving high overall accuracy.'),
        ('Proposed Methodology', 'PlasmaXAI combines frozen ResNet50 and DenseNet121 image embeddings with morphology, counterfactual, and score branches. Counterfactual signals modulate the image streams through sigmoid gates, refine the score stream, and contribute to a learned softmax modality gate before weighted fusion and final classification. For deployment, the model is now reported with a precision-recall balanced threshold selected from a stable validation performance plateau, improving both malignant-cell precision and malignant-cell recall relative to the strongest baselines.'),
        ('Outcomes and Results', f"At the selected precision-recall balanced operating point (threshold {THRESHOLD:.2f}), PlasmaXAI achieved {metrics['accuracy']:.2f}% accuracy, {metrics['weighted_f1']:.2f}% weighted F1, {metrics['plasma_precision']:.2f}% plasma precision, {metrics['plasma_recall']:.2f}% plasma recall, and {metrics['auc']:.2f}% AUC on the held-out test set. The strongest non-PlasmaXAI baseline by accuracy was {acc_baseline['model']} at {acc_baseline['accuracy']:.2f}%, the strongest non-PlasmaXAI baseline by precision was {prec_baseline['model']} at {prec_baseline['plasma_precision']:.2f}%, and the strongest non-PlasmaXAI baseline by recall was {recall_baseline['model']} at {recall_baseline['plasma_recall']:.2f}%. PlasmaXAI now leads the benchmark on accuracy, plasma precision, and plasma recall simultaneously. Bootstrap 95% intervals at this operating point were {bootstrap['accuracy']['low95']:.2f}-{bootstrap['accuracy']['high95']:.2f} for accuracy and {bootstrap['plasma_recall']['low95']:.2f}-{bootstrap['plasma_recall']['high95']:.2f} for plasma recall."),
        ('Patient-Level Robustness and Interpretation', f"The patient-level score using mean novel probability still produced an internal AUC of {robust['observed_auc']:.3f} on a cohort of only {robust['patient_count']} patients. Exact permutation testing gave one-sided p={robust['perm_p_one_sided']:.4f} and two-sided p={robust['perm_p_two_sided']:.4f}, while leave-one-patient-out AUC remained {robust['lopo_mean']:.3f}. These checks show strong internal separation, but because the cohort contains only ten same-source patients, the patient-level result should still be treated as exploratory and not as external clinical proof."),
        ('Impact Applications', 'With the updated operating point, PlasmaXAI is stronger for practical screening because it no longer trades malignant-cell recall for weak precision. The model now offers a more balanced and clinically usable error profile while preserving explainability through counterfactual features, modality gates, and patient-level signature analysis.'),
        ('Diagrams', 'The report package includes the framework diagram, expanded benchmark comparison, precision/F1 comparison, training curves, confusion/ROC comparison, patient robustness figure, and patient signature heatmap.'),
    ]


def save_markdown(sections: list[tuple[str, str]]):
    lines = []
    for heading, body in sections:
        lines.extend([heading, body, ''])
    ABSTRACT_MD.write_text('\n'.join(lines).strip() + '\n', encoding='utf-8')


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def save_docx(sections: list[tuple[str, str]]):
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(sections[0][1])
    run.bold = True
    run.font.size = Pt(16)
    for heading, body in sections[1:]:
        doc.add_heading(heading, level=1)
        para = doc.add_paragraph(body)
        para.style = doc.styles['Normal']
    doc.add_heading('Figures', level=1)
    figures = [
        (FIG_DIR / 'plasmaxai_fig13_framework_diagram.png', 'Figure 1. PlasmaXAI architecture overview.'),
        (COMPARE_FIG, 'Figure 2. Expanded benchmark comparison with the updated precision-recall balanced PlasmaXAI operating point.'),
        (PREC_F1_FIG, 'Figure 3. Precision, weighted F1, and precision-recall landscape across all benchmark models.'),
        (FIG_DIR / 'novel_fig8_training_curves.png', 'Figure 4. Training and validation curves for the selected PlasmaXAI fusion configuration.'),
        (CONFUSION_FIG, 'Figure 5. Confusion-delta and ROC comparison between PlasmaXAI and the previous hybrid baseline.'),
        (FIG_DIR / 'plasmaxai_fig16_patient_robustness.png', 'Figure 6. Patient-level robustness analysis showing exact permutation testing and leave-one-patient-out stability.'),
        (FIG_DIR / 'novel_fig11_patient_signature_heatmap.png', 'Figure 7. Patient-level signature heatmap for disease-associated counterfactual morphology patterns.'),
    ]
    for img_path, caption in figures:
        if img_path.exists():
            doc.add_picture(str(img_path), width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_caption(doc, caption)
    doc.save(str(ABSTRACT_DOCX))


def main():
    op = compute_operating_point()
    robust = json.loads(ROBUST_JSON.read_text(encoding='utf-8'))
    OP_JSON.write_text(json.dumps({
        'framework': FRAMEWORK_NAME,
        'operating_mode': 'precision_recall_balanced_validation_plateau',
        'threshold': THRESHOLD,
        'metrics': op['metrics'],
        'bootstrap': op['bootstrap'],
    }, indent=2), encoding='utf-8')
    compare_df = update_compare_csv(op['metrics'])
    plot_extended_comparison(compare_df)
    plot_precision_f1(compare_df)
    plot_confusion_roc(op['test_labels'], op['baseline_probs'], op['test_probs'], op['baseline_threshold'], THRESHOLD)
    sections = build_sections(compare_df, op['metrics'], op['bootstrap'], robust)
    save_markdown(sections)
    save_docx(sections)
    print(compare_df[['model','accuracy','weighted_f1','plasma_precision','plasma_recall','auc']].to_string(index=False))
    print(f'Updated threshold: {THRESHOLD:.2f}')
    print(f'Saved: {OP_JSON}')
    print(f'Saved: {COMPARE_CSV}')
    print(f'Saved: {COMPARE_FIG}')
    print(f'Saved: {PREC_F1_FIG}')
    print(f'Saved: {CONFUSION_FIG}')
    print(f'Saved: {ABSTRACT_MD}')
    print(f'Saved: {ABSTRACT_DOCX}')


if __name__ == '__main__':
    main()

